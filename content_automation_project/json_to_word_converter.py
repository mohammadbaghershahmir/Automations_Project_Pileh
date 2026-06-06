"""
JSON → Word conversion for Table Notes Generation output (ta*.json).

Headings only (no chapter, no italic): subchapter H1, topic H2, subtopic H3, subsubtopic H4.
Body from each row's ``points`` field.
"""

from __future__ import annotations

import json
import logging
import os
import re
from typing import Any, Dict, List, Optional, Tuple

try:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
except ImportError:  # pragma: no cover - exercised only when python-docx is missing
    OxmlElement = None  # type: ignore[misc, assignment]
    qn = None  # type: ignore[misc, assignment]

logger = logging.getLogger(__name__)

# Font that renders Persian and Latin parentheses reliably in Word.
DEFAULT_WORD_FONT = "Tahoma"

# Map decorative / full-width parentheses to ASCII so Word does not mirror them oddly.
_PAREN_NORMALIZATION = str.maketrans(
    {
        "\uFF08": "(",  # fullwidth left parenthesis
        "\uFF09": ")",  # fullwidth right parenthesis
        "\uFD3E": "(",  # ornate left parenthesis
        "\uFD3F": ")",  # ornate right parenthesis
    }
)

_BIDI_LANG_FA = "fa-IR"
_BIDI_LANG_EN = "en-US"

# Heading appearance without Word Heading styles (those force LTR/italic BiDi).
_HEADING_FONT_SIZES_PT: Dict[int, int] = {1: 16, 2: 14, 3: 12, 4: 11}

# English/Latin token islands in plain text (outside parentheses): "B6", "1.75".
_LTR_TOKEN_PATTERN = r"[A-Za-z0-9][A-Za-z0-9\s\-\./\+]*"
_LTR_TOKEN_RE = re.compile(_LTR_TOKEN_PATTERN)

# field_name → Word heading level (1–4)
HIERARCHY_HEADINGS: List[Tuple[str, int]] = [
    ("subchapter", 1),
    ("topic", 2),
    ("subtopic", 3),
    ("subsubtopic", 4),
]

HIERARCHY_FIELD_NAMES = [name for name, _ in HIERARCHY_HEADINGS]

_TA_FILENAME_RE = re.compile(r"^ta\d{6}_.+\.json$", re.IGNORECASE)


class TableNotesJsonError(ValueError):
    """Invalid or unsupported JSON for Table Notes → Word export."""


class DocumentProcessingJsonError(TableNotesJsonError):
    """Backward-compatible alias for runners that still import this name."""


def _normalize_key_map(row: Dict[str, Any]) -> Dict[str, Any]:
    return {str(k).lower(): v for k, v in row.items()}


def _row_field(row: Dict[str, Any], *names: str) -> str:
    lower = _normalize_key_map(row)
    for name in names:
        val = lower.get(name.lower())
        if val is not None and str(val).strip():
            return str(val).strip()
    return ""


def _looks_like_table_notes_filename(path_or_name: str) -> bool:
    base = os.path.basename(path_or_name or "")
    return bool(_TA_FILENAME_RE.match(base))


def validate_table_notes_json(
    data: Any,
    *,
    source_filename: str = "",
) -> Tuple[List[Dict[str, Any]], Optional[Dict[str, Any]]]:
    """
    Validate Table Notes (Stage TA) JSON and return the ``data`` row list.

    Requires top-level ``data`` (non-empty). Rejects bare Document Processing ``points`` files.
    """
    if not isinstance(data, dict):
        raise TableNotesJsonError("JSON root must be an object with a 'data' array.")

    if data.get("points") is not None and data.get("data") is None:
        raise TableNotesJsonError(
            "This file has top-level 'points' (Document Processing). "
            "Upload Table Notes output instead: ta{book}{chapter}_*.json with a 'data' array."
        )

    raw_rows = data.get("data")
    if raw_rows is None:
        stage = ""
        meta = data.get("metadata")
        if isinstance(meta, dict):
            stage = str(meta.get("stage") or "").strip()
        hint = f" (metadata.stage={stage!r})" if stage else ""
        raise TableNotesJsonError(
            "Missing top-level 'data' array. Upload Table Notes Generation output "
            f"(e.g. ta105003_Lesson_file_....json){hint}."
        )
    if not isinstance(raw_rows, list):
        raise TableNotesJsonError("'data' must be a JSON array.")
    if not raw_rows:
        raise TableNotesJsonError("'data' array is empty.")

    meta = data.get("metadata")
    metadata = meta if isinstance(meta, dict) else None
    if metadata is not None:
        stage = str(metadata.get("stage") or "").strip().upper()
        if stage and stage != "TA":
            raise TableNotesJsonError(
                f"Not Table Notes output (metadata.stage={metadata.get('stage')!r}). "
                "Upload ta{book}{chapter}_*.json from Table Notes Generation."
            )
        status = str(metadata.get("processing_status") or "").strip()
        if status and status != "completed":
            logger.warning(
                "Table Notes metadata.processing_status=%r (expected 'completed')",
                status,
            )
    elif source_filename and not _looks_like_table_notes_filename(source_filename):
        raise TableNotesJsonError(
            f"Filename {os.path.basename(source_filename)!r} does not match "
            "Table Notes pattern ta{book}{chapter}_*.json."
        )

    normalized: List[Dict[str, Any]] = []
    for i, item in enumerate(raw_rows):
        if not isinstance(item, dict):
            raise TableNotesJsonError(f"data[{i}] must be an object.")
        normalized.append(item)

    return normalized, metadata


def validate_document_processing_json(
    data: Any,
    *,
    source_filename: str = "",
) -> Tuple[List[Dict[str, Any]], Optional[Dict[str, Any]]]:
    """Alias: Table Notes validation (legacy import name)."""
    return validate_table_notes_json(data, source_filename=source_filename)


def _row_body(row: Dict[str, Any]) -> str:
    return _row_field(row, "points", "Points")


def _normalize_text_for_word(text: str) -> str:
    """Normalize punctuation that Word may render incorrectly in RTL paragraphs."""
    return text.translate(_PAREN_NORMALIZATION)


def _text_contains_rtl_script(text: str) -> bool:
    for char in text:
        if _char_script_direction(char) == "rtl":
            return True
    return False


def _char_script_direction(char: str) -> Optional[str]:
    code_point = ord(char)
    if (
        0x0590 <= code_point <= 0x05FF
        or 0x0600 <= code_point <= 0x06FF
        or 0x0750 <= code_point <= 0x077F
        or 0x08A0 <= code_point <= 0x08FF
        or 0xFB50 <= code_point <= 0xFDFF
        or 0xFE70 <= code_point <= 0xFEFF
    ):
        return "rtl"
    if ("A" <= char <= "Z") or ("a" <= char <= "z") or ("0" <= char <= "9"):
        return "ltr"
    return None


def _split_mixed_script(text: str) -> List[Tuple[str, bool]]:
    """Split text by script changes, keeping neutral punctuation with its neighbors."""
    if not text:
        return []

    segments: List[Tuple[str, bool]] = []
    buffer: List[str] = []
    buffer_rtl: Optional[bool] = None

    def flush() -> None:
        nonlocal buffer, buffer_rtl
        if not buffer:
            return
        is_rtl = buffer_rtl if buffer_rtl is not None else True
        segments.append(("".join(buffer), is_rtl))
        buffer = []
        buffer_rtl = None

    for char in text:
        direction = _char_script_direction(char)
        if direction is None:
            buffer.append(char)
            continue

        is_rtl = direction == "rtl"
        if buffer_rtl is None:
            buffer_rtl = is_rtl
            buffer.append(char)
            continue

        if is_rtl == buffer_rtl:
            buffer.append(char)
            continue

        flush()
        buffer_rtl = is_rtl
        buffer.append(char)

    flush()
    return segments


def _wrap_parenthetical(inner_segments: List[Tuple[str, bool]]) -> List[Tuple[str, bool]]:
    """Attach parentheses as separate pieces so ``)`` never sticks to an LTR run alone."""
    if not inner_segments:
        return [("()", True)]

    if len(inner_segments) == 1:
        segment_text, is_rtl = inner_segments[0]
        return [(f"({segment_text})", is_rtl)]

    wrapped: List[Tuple[str, bool]] = [("(", True)]
    wrapped.extend(inner_segments)
    wrapped.append((")", True))
    return wrapped


def _split_plain_ltr_islands(text: str) -> List[Tuple[str, bool]]:
    """Split non-parenthetical text into RTL chunks and LTR Latin/number islands."""
    if not text:
        return []

    segments: List[Tuple[str, bool]] = []
    cursor = 0
    for match in _LTR_TOKEN_RE.finditer(text):
        if match.start() > cursor:
            segments.append((text[cursor : match.start()], True))
        segments.append((match.group(0), False))
        cursor = match.end()

    if cursor < len(text):
        segments.append((text[cursor:], True))
    return segments


def _merge_adjacent_bidi_segments(segments: List[Tuple[str, bool]]) -> List[Tuple[str, bool]]:
    if not segments:
        return []

    merged: List[Tuple[str, bool]] = [segments[0]]
    for segment_text, is_rtl in segments[1:]:
        previous_text, previous_rtl = merged[-1]
        if is_rtl == previous_rtl:
            merged[-1] = (previous_text + segment_text, is_rtl)
        else:
            merged.append((segment_text, is_rtl))
    return merged


def _split_bidi_segments(text: str) -> List[Tuple[str, bool]]:
    """
    Split mixed Persian/English text into ordered (segment, is_rtl) pairs.

    Parentheses are handled explicitly:
    - ``(اتوزومال غالب)`` → one RTL run
    - ``(Family history)`` → one LTR run
    - ``(HIV سیفلیس و)`` / ``(مغلوب X مرتبط با)`` → split inside the parentheses
    """
    if not text:
        return []

    segments: List[Tuple[str, bool]] = []
    cursor = 0
    while cursor < len(text):
        if text[cursor] != "(":
            next_paren = text.find("(", cursor)
            chunk_end = len(text) if next_paren == -1 else next_paren
            segments.extend(_split_plain_ltr_islands(text[cursor:chunk_end]))
            cursor = chunk_end
            continue

        closing_paren = text.find(")", cursor + 1)
        if closing_paren == -1:
            segments.extend(_split_plain_ltr_islands(text[cursor:]))
            break

        inner_text = text[cursor + 1 : closing_paren]
        inner_segments = _split_mixed_script(inner_text)
        segments.extend(_wrap_parenthetical(inner_segments))
        cursor = closing_paren + 1

    return _merge_adjacent_bidi_segments(segments)


def _add_bidi_run(
    paragraph: Any,
    segment_text: str,
    *,
    is_rtl: bool,
    font_name: str,
    size_pt: Optional[int],
    bold: bool,
) -> None:
    """Append one run with plain text (no Unicode control chars) and correct OOXML direction."""
    run = paragraph.add_run(segment_text)
    _apply_run_font(run, font_name, size_pt=size_pt)
    _set_run_rtl(run, is_rtl=is_rtl)
    _set_run_language(run, is_rtl=is_rtl)
    run.bold = bold
    run.italic = False


def _set_paragraph_jc_right(paragraph: Any) -> None:
    """Set physical right alignment (same as Word's Right Align toolbar button)."""
    if OxmlElement is None or qn is None:
        return

    paragraph_properties = paragraph._element.get_or_add_pPr()
    justification = paragraph_properties.find(qn("w:jc"))
    if justification is None:
        justification = OxmlElement("w:jc")
        paragraph_properties.insert_element_before(
            justification,
            "w:rPr",
            "w:sectPr",
            "w:pPrChange",
        )
    justification.set(qn("w:val"), "right")

    try:
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    except ImportError:
        pass


def _set_paragraph_rtl(paragraph: Any) -> None:
    """Mark a paragraph RTL (``w:bidi``) and right-align it."""
    if OxmlElement is None or qn is None:
        return

    paragraph_properties = paragraph._element.get_or_add_pPr()
    if paragraph_properties.find(qn("w:bidi")) is None:
        bidi = OxmlElement("w:bidi")
        paragraph_properties.insert_element_before(
            bidi,
            "w:jc",
            "w:rPr",
            "w:sectPr",
            "w:pPrChange",
        )

    _set_paragraph_jc_right(paragraph)

    default_run = paragraph_properties.find(qn("w:rPr"))
    if default_run is None:
        default_run = OxmlElement("w:rPr")
        paragraph_properties.append(default_run)
    if default_run.find(qn("w:rtl")) is None:
        default_run.append(OxmlElement("w:rtl"))
    default_lang = default_run.find(qn("w:lang"))
    if default_lang is None:
        default_lang = OxmlElement("w:lang")
        default_run.append(default_lang)
    default_lang.set(qn("w:bidi"), _BIDI_LANG_FA)
    default_lang.set(qn("w:val"), _BIDI_LANG_FA)


def _apply_run_font(run: Any, font_name: str = DEFAULT_WORD_FONT, *, size_pt: Optional[int] = None) -> None:
    """Apply font to Latin and complex-script (Persian) text in a run."""
    run.font.name = font_name
    if size_pt is not None:
        try:
            from docx.shared import Pt

            run.font.size = Pt(size_pt)
        except ImportError:
            pass

    if qn is None:
        return

    run_properties = run._element.get_or_add_rPr()
    font_element = run_properties.get_or_add_rFonts()
    font_element.set(qn("w:ascii"), font_name)
    font_element.set(qn("w:hAnsi"), font_name)
    font_element.set(qn("w:cs"), font_name)
    font_element.set(qn("w:hint"), "cs")

    if run_properties.find(qn("w:cs")) is None:
        run_properties.append(OxmlElement("w:cs"))

    if size_pt is not None:
        size_value = str(int(size_pt * 2))
        for tag in ("w:sz", "w:szCs"):
            size_element = run_properties.find(qn(tag))
            if size_element is None:
                size_element = OxmlElement(tag)
                run_properties.append(size_element)
            size_element.set(qn("w:val"), size_value)


def _set_run_rtl(run: Any, *, is_rtl: bool) -> None:
    """Toggle ``w:rtl`` on a run. LTR islands (English) must not carry this flag."""
    if OxmlElement is None or qn is None:
        return

    run_properties = run._element.get_or_add_rPr()
    rtl_element = run_properties.find(qn("w:rtl"))
    if is_rtl:
        if rtl_element is None:
            run_properties.append(OxmlElement("w:rtl"))
        return

    if rtl_element is not None:
        run_properties.remove(rtl_element)


def _set_run_language(run: Any, *, is_rtl: bool) -> None:
    """Set complex-script / Latin language tags on a run."""
    if OxmlElement is None or qn is None:
        return

    run_properties = run._element.get_or_add_rPr()
    for existing in run_properties.findall(qn("w:lang")):
        run_properties.remove(existing)

    language = OxmlElement("w:lang")
    if is_rtl:
        language.set(qn("w:bidi"), _BIDI_LANG_FA)
        language.set(qn("w:val"), _BIDI_LANG_FA)
    else:
        language.set(qn("w:val"), _BIDI_LANG_EN)
    run_properties.append(language)


def _clear_paragraph_runs(paragraph: Any) -> None:
    for run in list(paragraph.runs):
        paragraph._element.remove(run._element)


def _populate_paragraph_with_bidi_text(
    paragraph: Any,
    text: str,
    *,
    font_name: str = DEFAULT_WORD_FONT,
    size_pt: Optional[int] = None,
    bold: bool = False,
) -> None:
    """
    Fill a paragraph with BiDi-aware Persian text.

    Uses separate OOXML runs (``w:rtl`` on/off per segment). Do not embed Unicode
    BiDi control characters — Word may show them as visible ``RLI``/``PDI`` boxes.
    """
    normalized_text = _normalize_text_for_word(text)
    _clear_paragraph_runs(paragraph)

    if not normalized_text.strip():
        run = paragraph.add_run(normalized_text)
        _apply_run_font(run, font_name, size_pt=size_pt)
        _set_paragraph_jc_right(paragraph)
        return

    if not _text_contains_rtl_script(normalized_text):
        run = paragraph.add_run(normalized_text)
        _apply_run_font(run, font_name, size_pt=size_pt)
        run.bold = bold
        run.italic = False
        _set_paragraph_jc_right(paragraph)
        return

    _set_paragraph_rtl(paragraph)
    segments = _split_bidi_segments(normalized_text)
    if not segments:
        segments = [(normalized_text, True)]

    for segment_text, is_rtl in segments:
        if not segment_text:
            continue
        _add_bidi_run(
            paragraph,
            segment_text,
            is_rtl=is_rtl,
            font_name=font_name,
            size_pt=size_pt,
            bold=bold,
        )


def _add_bidi_paragraph(doc: Any, text: str) -> Any:
    paragraph = doc.add_paragraph()
    _populate_paragraph_with_bidi_text(paragraph, text)
    return paragraph


def _configure_document_defaults(doc: Any) -> None:
    """Configure Normal style for Persian RTL (avoid Word Heading styles for BiDi)."""
    try:
        normal = doc.styles["Normal"]
    except KeyError:
        return

    normal.font.name = DEFAULT_WORD_FONT
    normal.font.italic = False

    if OxmlElement is None or qn is None:
        return

    style_element = normal._element
    style_paragraph = style_element.find(qn("w:pPr"))
    if style_paragraph is None:
        style_paragraph = OxmlElement("w:pPr")
        style_element.insert(0, style_paragraph)
    if style_paragraph.find(qn("w:bidi")) is None:
        style_paragraph.append(OxmlElement("w:bidi"))

    style_run = style_paragraph.find(qn("w:rPr"))
    if style_run is None:
        style_run = OxmlElement("w:rPr")
        style_paragraph.append(style_run)
    if style_run.find(qn("w:rtl")) is None:
        style_run.append(OxmlElement("w:rtl"))
    if style_run.find(qn("w:cs")) is None:
        style_run.append(OxmlElement("w:cs"))

    style_fonts = style_run.find(qn("w:rFonts"))
    if style_fonts is None:
        style_fonts = OxmlElement("w:rFonts")
        style_run.insert(0, style_fonts)
    style_fonts.set(qn("w:cs"), DEFAULT_WORD_FONT)

    style_lang = style_run.find(qn("w:lang"))
    if style_lang is None:
        style_lang = OxmlElement("w:lang")
        style_run.append(style_lang)
    style_lang.set(qn("w:bidi"), _BIDI_LANG_FA)
    style_lang.set(qn("w:val"), _BIDI_LANG_FA)

    style_jc = style_paragraph.find(qn("w:jc"))
    if style_jc is None:
        style_jc = OxmlElement("w:jc")
        style_paragraph.append(style_jc)
    style_jc.set(qn("w:val"), "right")

    try:
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

        normal.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    except ImportError:
        pass


def _disable_heading_italic_styles(doc: Any) -> None:
    """Prepare document defaults; headings use Normal + manual bold/size (not Heading 1–4)."""
    _configure_document_defaults(doc)


def _add_heading_no_italic(doc: Any, text: str, level: int) -> None:
    """Add a hierarchy heading without Word Heading styles (they break Persian BiDi)."""
    paragraph = doc.add_paragraph(style="Normal")
    size_pt = _HEADING_FONT_SIZES_PT.get(level, 11)
    _populate_paragraph_with_bidi_text(
        paragraph,
        text,
        size_pt=size_pt,
        bold=True,
    )


def convert_points_to_docx(points: List[Dict[str, Any]], output_path: str) -> bool:
    """Build a .docx from Table Notes ``data`` rows."""
    try:
        from docx import Document
    except ImportError as e:
        logger.error("python-docx is required for JSON to Word: %s", e)
        return False

    doc = Document()
    _disable_heading_italic_styles(doc)
    last: Dict[str, str] = {name: "" for name in HIERARCHY_FIELD_NAMES}
    paragraphs_added = 0

    for row in points:
        for field, level in HIERARCHY_HEADINGS:
            value = _row_field(row, field)
            if not value or value == last[field]:
                continue
            _add_heading_no_italic(doc, value, level)
            last[field] = value
            idx = HIERARCHY_FIELD_NAMES.index(field)
            for reset_field in HIERARCHY_FIELD_NAMES[idx + 1 :]:
                last[reset_field] = ""

        body = _row_body(row)
        if not body:
            continue
        _add_bidi_paragraph(doc, body)
        paragraphs_added += 1

    if paragraphs_added == 0:
        logger.error("No point body text found in data rows")
        return False

    os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
    doc.save(output_path)
    logger.info("Wrote Word document: %s (%d body paragraphs)", output_path, paragraphs_added)
    return True


def convert_json_file_to_docx(json_path: str, docx_path: str) -> bool:
    """Read a Table Notes JSON file and write a .docx."""
    try:
        with open(json_path, "r", encoding="utf-8") as f:
            data = json.load(f)
    except json.JSONDecodeError as e:
        logger.error("Invalid JSON in %s: %s", json_path, e)
        return False
    except OSError as e:
        logger.error("Cannot read %s: %s", json_path, e)
        return False

    try:
        rows, _meta = validate_table_notes_json(data, source_filename=json_path)
    except TableNotesJsonError as e:
        logger.error("%s", e)
        return False

    return convert_points_to_docx(rows, docx_path)
