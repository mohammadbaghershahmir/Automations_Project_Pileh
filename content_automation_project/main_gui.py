"""
Main GUI for Content Automation Project
Part 1: PDF Upload, Prompt Selection, Model Selection, and AI Studio Integration
"""

import tkinter as tk
from tkinter import filedialog, messagebox, ttk
import customtkinter as ctk
import os
import logging
import threading
import csv
import io
import json
import glob
from datetime import datetime
from typing import Optional, List, Dict, Any

from api_layer import APIConfig, APIKeyManager, GeminiAPIClient
from pdf_processor import PDFProcessor
from prompt_manager import PromptManager
from multi_part_processor import MultiPartProcessor
from multi_part_post_processor import MultiPartPostProcessor


class ContentAutomationGUI:
    """Main GUI application for content automation"""
    
    def __init__(self):
        # Configure appearance
        ctk.set_appearance_mode("dark")
        ctk.set_default_color_theme("blue")
        
        # Initialize main window
        self.root = ctk.CTk()
        self.root.title("Content Automation - Part 1")
        self.root.geometry("1000x800")
        self.root.minsize(900, 700)

        # Common Farsi-friendly font for textboxes
        try:
            self.farsi_text_font = ctk.CTkFont(family="Tahoma", size=11)
        except Exception:
            # Fallback if Tahoma is not available
            self.farsi_text_font = ctk.CTkFont(size=11)
        
        # Setup logging
        self.setup_logging()
        
        # Initialize components
        self.pdf_processor = PDFProcessor()
        self.prompt_manager = PromptManager()
        self.api_key_manager = APIKeyManager()
        self.api_client = GeminiAPIClient(self.api_key_manager)
        self.multi_part_processor = MultiPartProcessor(self.api_client)
        self.multi_part_post_processor = MultiPartPostProcessor(self.api_client)
        
        # Variables
        self.pdf_path = None
        self.selected_prompt_name = None
        self.custom_prompt = ""
        self.use_custom_prompt = False
        self.last_final_output_path = None       # Stage 1 JSON
        self.last_post_processed_path = None     # Stage 2 JSON
        self.last_corrected_path = None          # Stage 3 JSON
        
        # Setup UI
        self.setup_ui()
        
    def setup_logging(self):
        """Setup logging configuration"""
        logging.basicConfig(
            level=logging.INFO,
            format='%(asctime)s - %(levelname)s - %(message)s',
            handlers=[
                logging.FileHandler('content_automation.log', encoding='utf-8'),
                logging.StreamHandler()
            ]
        )
        self.logger = logging.getLogger(__name__)
    
    def setup_ui(self):
        """Setup the main user interface"""
        # Create main container with scrollable frame
        main_frame = ctk.CTkScrollableFrame(self.root)
        main_frame.pack(fill="both", expand=True, padx=20, pady=20)
        
        # Title
        title = ctk.CTkLabel(main_frame, text="Content Automation - Part 1", 
                            font=ctk.CTkFont(size=24, weight="bold"))
        title.pack(pady=(0, 20))
        
        # API Configuration Section
        self.setup_api_section(main_frame)
        
        # PDF Upload Section
        self.setup_pdf_section(main_frame)
        
        # Prompt Selection Section
        self.setup_prompt_section(main_frame)
        
        # Model Selection Section
        self.setup_model_section(main_frame)
        
        # Output Section
        self.setup_output_section(main_frame)
        
        # Control Buttons
        self.setup_controls(main_frame)
        
        # Status Section
        self.setup_status_section(main_frame)
    
    def setup_api_section(self, parent):
        """Setup API configuration section"""
        api_frame = ctk.CTkFrame(parent)
        api_frame.pack(fill="x", pady=(0, 20))
        
        api_label = ctk.CTkLabel(api_frame, text="🔑 API Configuration", 
                                font=ctk.CTkFont(size=18, weight="bold"))
        api_label.pack(pady=(15, 10))
        
        # API Key CSV file selection
        key_frame = ctk.CTkFrame(api_frame)
        key_frame.pack(fill="x", padx=15, pady=5)
        
        ctk.CTkLabel(key_frame, text="API Key CSV File:", 
                    font=ctk.CTkFont(size=12, weight="bold")).pack(anchor="w", padx=10, pady=(10, 5))
        
        self.api_key_file_var = ctk.StringVar()
        api_key_entry = ctk.CTkEntry(key_frame, textvariable=self.api_key_file_var, width=400)
        api_key_entry.pack(side="left", fill="x", expand=True, padx=(10, 5), pady=(0, 5))
        
        ctk.CTkButton(key_frame, text="Browse", command=self.browse_api_key_file, 
                     width=80).pack(side="right", padx=(5, 10), pady=(0, 5))
        
        ctk.CTkLabel(key_frame, text="CSV format: account;project;api_key (API keys will be used in rotation)", 
                    font=ctk.CTkFont(size=10), text_color="gray").pack(anchor="w", padx=10, pady=(0, 10))
        
        # API keys status
        self.api_keys_status_label = ctk.CTkLabel(key_frame, text="No API keys loaded", 
                                                  font=ctk.CTkFont(size=10), text_color="gray")
        self.api_keys_status_label.pack(anchor="w", padx=10, pady=(0, 10))
    
    def setup_pdf_section(self, parent):
        """Setup PDF upload section"""
        pdf_frame = ctk.CTkFrame(parent)
        pdf_frame.pack(fill="x", pady=(0, 20))
        
        pdf_label = ctk.CTkLabel(pdf_frame, text="📄 PDF Upload", 
                                font=ctk.CTkFont(size=18, weight="bold"))
        pdf_label.pack(pady=(15, 10))
        
        # PDF file selection
        file_frame = ctk.CTkFrame(pdf_frame)
        file_frame.pack(fill="x", padx=15, pady=5)
        
        ctk.CTkLabel(file_frame, text="PDF File:", 
                    font=ctk.CTkFont(size=12, weight="bold")).pack(anchor="w", padx=10, pady=(10, 5))
        
        self.pdf_file_var = ctk.StringVar()
        pdf_entry = ctk.CTkEntry(file_frame, textvariable=self.pdf_file_var, width=400)
        pdf_entry.pack(side="left", fill="x", expand=True, padx=(10, 5), pady=(0, 5))
        
        ctk.CTkButton(file_frame, text="Browse", command=self.browse_pdf_file, 
                     width=80).pack(side="right", padx=(5, 10), pady=(0, 5))
        
        # PDF info display
        self.pdf_info_label = ctk.CTkLabel(file_frame, text="No PDF selected", 
                                           font=ctk.CTkFont(size=10), text_color="gray")
        self.pdf_info_label.pack(anchor="w", padx=10, pady=(0, 10))
    
    def setup_prompt_section(self, parent):
        """Setup prompt selection section"""
        prompt_frame = ctk.CTkFrame(parent)
        prompt_frame.pack(fill="x", pady=(0, 20))
        
        prompt_label = ctk.CTkLabel(prompt_frame, text="💬 Prompt Selection", 
                                   font=ctk.CTkFont(size=18, weight="bold"))
        prompt_label.pack(pady=(15, 10))
        
        # Prompt type selection
        type_frame = ctk.CTkFrame(prompt_frame)
        type_frame.pack(fill="x", padx=15, pady=5)
        
        self.prompt_type_var = ctk.StringVar(value="predefined")
        ctk.CTkRadioButton(type_frame, text="Use Predefined Prompt", 
                          variable=self.prompt_type_var, value="predefined",
                          command=self.on_prompt_type_change).pack(anchor="w", padx=10, pady=5)
        ctk.CTkRadioButton(type_frame, text="Use Custom Prompt", 
                          variable=self.prompt_type_var, value="custom",
                          command=self.on_prompt_type_change).pack(anchor="w", padx=10, pady=5)
        
        # Predefined prompt selection
        predefined_frame = ctk.CTkFrame(prompt_frame)
        predefined_frame.pack(fill="x", padx=15, pady=5)
        
        ctk.CTkLabel(predefined_frame, text="Select Predefined Prompt:", 
                    font=ctk.CTkFont(size=12, weight="bold")).pack(anchor="w", padx=10, pady=(10, 5))
        
        prompt_names = self.prompt_manager.get_prompt_names()
        default_value = prompt_names[0] if prompt_names else ""
        self.prompt_combo_var = ctk.StringVar(value=default_value)
        self.prompt_combo = ctk.CTkComboBox(predefined_frame, values=prompt_names, 
                                           variable=self.prompt_combo_var, width=400,
                                           command=self.on_prompt_selected)
        self.prompt_combo.pack(anchor="w", padx=10, pady=(0, 5))
        
        # Prompt preview
        self.prompt_preview = ctk.CTkTextbox(predefined_frame, height=100, font=self.farsi_text_font)
        self.prompt_preview.pack(fill="x", padx=10, pady=(0, 5))
        if prompt_names and default_value:
            preview_text = self.prompt_manager.get_prompt(default_value) or ""
            self.prompt_preview.insert("1.0", preview_text)
            self.selected_prompt_name = default_value
        self.prompt_preview.configure(state="disabled")
        
        # Custom prompt input
        self.custom_frame = ctk.CTkFrame(prompt_frame)
        
        ctk.CTkLabel(self.custom_frame, text="Enter Custom Prompt:", 
                    font=ctk.CTkFont(size=12, weight="bold")).pack(anchor="w", padx=10, pady=(10, 5))
        
        self.custom_prompt_text = ctk.CTkTextbox(self.custom_frame, height=120, font=self.farsi_text_font)
        self.custom_prompt_text.pack(fill="x", padx=10, pady=(0, 10))
        
        # Update visibility based on initial selection
        self.on_prompt_type_change()
    
    def setup_model_section(self, parent):
        """Setup model selection section"""
        model_frame = ctk.CTkFrame(parent)
        model_frame.pack(fill="x", pady=(0, 20))
        
        model_label = ctk.CTkLabel(model_frame, text="🤖 Model Selection", 
                                  font=ctk.CTkFont(size=18, weight="bold"))
        model_label.pack(pady=(15, 10))
        
        model_select_frame = ctk.CTkFrame(model_frame)
        model_select_frame.pack(fill="x", padx=15, pady=(0, 15))
        
        ctk.CTkLabel(model_select_frame, text="Select Gemini Model:", 
                    font=ctk.CTkFont(size=12, weight="bold")).pack(anchor="w", padx=10, pady=(10, 5))
        
        self.model_var = ctk.StringVar(value=APIConfig.TEXT_MODELS[0])
        self.model_combo = ctk.CTkComboBox(model_select_frame, values=APIConfig.TEXT_MODELS, 
                                          variable=self.model_var, width=400)
        self.model_combo.pack(anchor="w", padx=10, pady=(0, 10))
        
        ctk.CTkLabel(model_select_frame, 
                    text="Available models: gemini-2.5-flash, gemini-2.5-pro, gemini-2.0-flash, etc.", 
                    font=ctk.CTkFont(size=10), text_color="gray").pack(anchor="w", padx=10, pady=(0, 10))
    
    def setup_output_section(self, parent):
        """Setup output configuration section"""
        output_frame = ctk.CTkFrame(parent)
        output_frame.pack(fill="x", pady=(0, 20))
        
        output_label = ctk.CTkLabel(output_frame, text="💾 Output Configuration", 
                                   font=ctk.CTkFont(size=18, weight="bold"))
        output_label.pack(pady=(15, 10))
        
        # Output folder selection
        folder_frame = ctk.CTkFrame(output_frame)
        folder_frame.pack(fill="x", padx=15, pady=5)
        
        ctk.CTkLabel(folder_frame, text="Output Folder:", 
                    font=ctk.CTkFont(size=12, weight="bold")).pack(anchor="w", padx=10, pady=(10, 5))
        
        self.output_folder_var = ctk.StringVar()
        output_entry = ctk.CTkEntry(folder_frame, textvariable=self.output_folder_var, width=400)
        output_entry.pack(side="left", fill="x", expand=True, padx=(10, 5), pady=(0, 5))
        
        ctk.CTkButton(folder_frame, text="Browse", command=self.browse_output_folder, 
                     width=80).pack(side="right", padx=(5, 10), pady=(0, 5))
        
        ctk.CTkLabel(folder_frame, text="Responses will be saved as text files in this folder", 
                    font=ctk.CTkFont(size=10), text_color="gray").pack(anchor="w", padx=10, pady=(0, 10))
    
    def setup_controls(self, parent):
        """Setup control buttons"""
        controls_frame = ctk.CTkFrame(parent)
        controls_frame.pack(fill="x", pady=(0, 20))
        
        buttons_frame = ctk.CTkFrame(controls_frame)
        buttons_frame.pack(pady=20)
        
        # Process button
        self.process_btn = ctk.CTkButton(buttons_frame, text="🚀 Process PDF with AI", 
                                        command=self.process_pdf, width=200, height=40,
                                        font=ctk.CTkFont(size=14, weight="bold"))
        self.process_btn.pack(side="left", padx=10)
        
        # View CSV button
        self.view_csv_btn = ctk.CTkButton(buttons_frame, text="📊 View CSV", 
                                         command=self.view_csv_from_json, width=150, height=40,
                                         font=ctk.CTkFont(size=14),
                                         fg_color="green", hover_color="darkgreen")
        self.view_csv_btn.pack(side="left", padx=10)

        # View/Edit Stage 1 JSON button
        self.view_stage1_json_btn = ctk.CTkButton(
            buttons_frame,
            text="📄 View/Edit Stage 1 JSON",
            command=lambda: self.open_json_editor(self.last_final_output_path),
            width=220,
            height=40,
            font=ctk.CTkFont(size=14),
        )
        self.view_stage1_json_btn.pack(side="left", padx=10)

        # Next step button - go to second page (post-processing of final JSON by Part)
        self.next_step_btn = ctk.CTkButton(
            buttons_frame,
            text="Next Step (Part Processing)",
            command=self.open_part_processing_window,
            width=220,
            height=40,
            font=ctk.CTkFont(size=14, weight="bold"),
        )
        self.next_step_btn.pack(side="left", padx=10)
    
    def setup_status_section(self, parent):
        """Setup status section"""
        status_frame = ctk.CTkFrame(parent)
        status_frame.pack(fill="x", pady=(0, 20))
        
        ctk.CTkLabel(status_frame, text="📊 Status", 
                    font=ctk.CTkFont(size=18, weight="bold")).pack(pady=(15, 10))
        
        self.status_text = ctk.CTkTextbox(status_frame, height=150, font=self.farsi_text_font)
        self.status_text.pack(fill="x", padx=15, pady=(0, 15))
        
        self.update_status("Ready. Please configure API keys and select a PDF file.")
    
    def browse_api_key_file(self):
        """Browse for API key CSV file"""
        filename = filedialog.askopenfilename(
            title="Select API Key CSV File",
            filetypes=[("CSV files", "*.csv"), ("All files", "*.*")]
        )
        if filename:
            self.api_key_file_var.set(filename)
            if self.api_key_manager.load_from_csv(filename):
                num_keys = len(self.api_key_manager.api_keys)
                self.api_keys_status_label.configure(
                    text=f"✓ Loaded {num_keys} API key(s) - Will be used in rotation",
                    text_color="green"
                )
                self.update_status(f"Loaded {num_keys} API key(s) from: {os.path.basename(filename)}")
            else:
                self.api_keys_status_label.configure(
                    text="✗ Failed to load API keys",
                    text_color="red"
                )
                messagebox.showerror("Error", "Failed to load API keys from file")
    
    
    def browse_pdf_file(self):
        """Browse for PDF file"""
        filename = filedialog.askopenfilename(
            title="Select PDF File",
            filetypes=[("PDF files", "*.pdf"), ("All files", "*.*")]
        )
        if filename:
            self.pdf_file_var.set(filename)
            self.validate_pdf(filename)
    
    def browse_output_folder(self):
        """Browse for output folder"""
        folder = filedialog.askdirectory(title="Select Output Folder")
        if folder:
            self.output_folder_var.set(folder)
            self.update_status(f"Output folder set: {folder}")
    
    def validate_pdf(self, file_path: str):
        """Validate PDF file and update UI"""
        is_valid, error_msg, page_count = self.pdf_processor.validate_pdf(file_path)
        
        if is_valid:
            info = self.pdf_processor.get_pdf_info(file_path)
            file_size_mb = info['file_size'] / (1024 * 1024)
            self.pdf_info_label.configure(
                text=f"✓ Valid PDF - {page_count} pages, {file_size_mb:.2f} MB",
                text_color="green"
            )
            self.pdf_path = file_path
            self.update_status(f"PDF validated: {os.path.basename(file_path)} ({page_count} pages)")
        else:
            self.pdf_info_label.configure(
                text=f"✗ {error_msg}",
                text_color="red"
            )
            self.pdf_path = None
            messagebox.showerror("PDF Validation Error", error_msg)
    
    def on_prompt_type_change(self):
        """Handle prompt type change"""
        if self.prompt_type_var.get() == "predefined":
            # Show predefined, hide custom
            predefined_frame = self.prompt_combo.master
            self.prompt_combo.pack(anchor="w", padx=10, pady=(0, 5))
            self.prompt_preview.pack(fill="x", padx=10, pady=(0, 5))
            self.custom_frame.pack_forget()
            self.use_custom_prompt = False
        else:
            # Hide predefined, show custom
            predefined_frame = self.prompt_combo.master
            self.prompt_combo.pack_forget()
            self.prompt_preview.pack_forget()
            self.custom_frame.pack(fill="x", padx=15, pady=5)
            self.use_custom_prompt = True
    
    def on_prompt_selected(self, selected_name: str):
        """Handle predefined prompt selection"""
        prompt_text = self.prompt_manager.get_prompt(selected_name)
        if prompt_text:
            self.prompt_preview.configure(state="normal")
            self.prompt_preview.delete("1.0", tk.END)
            self.prompt_preview.insert("1.0", prompt_text)
            self.prompt_preview.configure(state="disabled")
            self.selected_prompt_name = selected_name
    
    def get_selected_prompt(self) -> Optional[str]:
        """Get the selected prompt text"""
        if self.use_custom_prompt:
            prompt = self.custom_prompt_text.get("1.0", tk.END).strip()
            return prompt if prompt else None
        else:
            return self.prompt_manager.get_prompt(self.prompt_combo_var.get())
    
    def update_status(self, message: str):
        """Update status text"""
        self.status_text.delete("1.0", tk.END)
        self.status_text.insert("1.0", f"[{self.get_timestamp()}] {message}")
    
    def get_timestamp(self) -> str:
        """Get current timestamp"""
        from datetime import datetime
        return datetime.now().strftime("%H:%M:%S")

    def open_part_processing_window(self):
        """
        Open second page for part-by-part processing of an existing final_output.json.
        """
        window = ctk.CTkToplevel(self.root)
        window.title("Part Processing - Second Stage")
        window.geometry("900x700")

        main_frame = ctk.CTkScrollableFrame(window)
        main_frame.pack(fill="both", expand=True, padx=20, pady=20)

        title = ctk.CTkLabel(
            main_frame,
            text="Second Stage - Process JSON by Part",
            font=ctk.CTkFont(size=22, weight="bold"),
        )
        title.pack(pady=(0, 20))

        # Prompt section (similar to first form but simplified)
        prompt_frame = ctk.CTkFrame(main_frame)
        prompt_frame.pack(fill="x", pady=(0, 20))

        ctk.CTkLabel(
            prompt_frame,
            text="💬 Second-Stage Prompt",
            font=ctk.CTkFont(size=18, weight="bold"),
        ).pack(pady=(15, 10))

        ctk.CTkLabel(
            prompt_frame,
            text="Enter the prompt that should be applied separately to each Part.\n"
                 "The JSON rows for each Part will be sent to the model with this prompt.",
            font=ctk.CTkFont(size=11),
            text_color="gray",
        ).pack(anchor="w", padx=10, pady=(0, 10))

        self.second_stage_prompt_text = ctk.CTkTextbox(prompt_frame, height=140, font=self.farsi_text_font)
        self.second_stage_prompt_text.pack(fill="x", padx=10, pady=(0, 10))

        # Pre-fill with current prompt if available
        current_prompt = self.get_selected_prompt()
        if current_prompt:
            self.second_stage_prompt_text.insert("1.0", current_prompt)

        # Chapter name section (right after prompt)
        chapter_frame = ctk.CTkFrame(prompt_frame)
        chapter_frame.pack(fill="x", padx=10, pady=(0, 10))

        ctk.CTkLabel(
            chapter_frame,
            text="Chapter Name:",
            font=ctk.CTkFont(size=12, weight="bold"),
        ).pack(anchor="w", padx=10, pady=(10, 5))

        self.second_stage_chapter_var = ctk.StringVar(value="درمان با UV")
        chapter_entry = ctk.CTkEntry(
            chapter_frame,
            textvariable=self.second_stage_chapter_var,
            width=400
        )
        chapter_entry.pack(anchor="w", padx=10, pady=(0, 5))

        ctk.CTkLabel(
            chapter_frame,
            text="This name will automatically replace {CHAPTER_NAME} in your prompt.",
            font=ctk.CTkFont(size=10),
            text_color="gray",
        ).pack(anchor="w", padx=10, pady=(0, 10))

        # JSON selection section
        json_frame = ctk.CTkFrame(main_frame)
        json_frame.pack(fill="x", pady=(0, 20))

        ctk.CTkLabel(
            json_frame,
            text="📄 Input JSON (final_output.json from previous step)",
            font=ctk.CTkFont(size=18, weight="bold"),
        ).pack(pady=(15, 10))

        file_frame = ctk.CTkFrame(json_frame)
        file_frame.pack(fill="x", padx=15, pady=5)

        ctk.CTkLabel(
            file_frame,
            text="JSON File:",
            font=ctk.CTkFont(size=12, weight="bold"),
        ).pack(anchor="w", padx=10, pady=(10, 5))

        self.second_stage_json_var = ctk.StringVar()
        if self.last_final_output_path and os.path.exists(self.last_final_output_path):
            self.second_stage_json_var.set(self.last_final_output_path)

        json_entry = ctk.CTkEntry(file_frame, textvariable=self.second_stage_json_var, width=400)
        json_entry.pack(side="left", fill="x", expand=True, padx=(10, 5), pady=(0, 5))

        def browse_json_file():
            filename = filedialog.askopenfilename(
                title="Select final_output.json file",
                filetypes=[("JSON files", "*.json"), ("All files", "*.*")]
            )
            if filename:
                self.second_stage_json_var.set(filename)

        ctk.CTkButton(
            file_frame,
            text="Browse",
            command=browse_json_file,
            width=80,
        ).pack(side="right", padx=(5, 10), pady=(0, 5))

        # Model selection for second stage
        model_frame = ctk.CTkFrame(main_frame)
        model_frame.pack(fill="x", pady=(0, 20))

        ctk.CTkLabel(
            model_frame,
            text="🤖 Second-Stage Model Selection",
            font=ctk.CTkFont(size=18, weight="bold"),
        ).pack(pady=(15, 10))

        inner_model_frame = ctk.CTkFrame(model_frame)
        inner_model_frame.pack(fill="x", padx=15, pady=(0, 10))

        ctk.CTkLabel(
            inner_model_frame,
            text="Select model for second-stage (per-Part) processing:",
            font=ctk.CTkFont(size=12, weight="bold"),
        ).pack(anchor="w", padx=10, pady=(10, 5))

        # Default to the model selected in the first page
        self.second_stage_model_var = ctk.StringVar(value=self.model_var.get())
        self.second_stage_model_combo = ctk.CTkComboBox(
            inner_model_frame,
            values=APIConfig.TEXT_MODELS,
            variable=self.second_stage_model_var,
            width=400,
        )
        self.second_stage_model_combo.pack(anchor="w", padx=10, pady=(0, 10))

        ctk.CTkLabel(
            inner_model_frame,
            text="You can use a lighter model (e.g., gemini-2.5-flash) for faster processing in the second stage.",
            font=ctk.CTkFont(size=10),
            text_color="gray",
        ).pack(anchor="w", padx=10, pady=(0, 5))

        # Control buttons in second page
        controls_frame = ctk.CTkFrame(main_frame)
        controls_frame.pack(fill="x", pady=(10, 10))

        def start_second_stage():
            # Disable button to prevent multiple clicks during processing
            start_button.configure(state="disabled", text="Processing...")
            threading.Thread(
                target=self.process_json_by_parts_worker,
                args=(window, start_button),
                daemon=True
            ).start()

        start_button = ctk.CTkButton(
            controls_frame,
            text="🚀 Start Part Processing",
            command=start_second_stage,
            width=220,
            height=40,
            font=ctk.CTkFont(size=14, weight="bold"),
        )
        start_button.pack(side="left", padx=10, pady=10)

        # View/Edit Stage 2 JSON button
        ctk.CTkButton(
            controls_frame,
            text="View/Edit Stage 2 JSON",
            command=lambda: self.open_json_editor(self.last_post_processed_path),
            width=200,
            height=40,
            font=ctk.CTkFont(size=13),
        ).pack(side="left", padx=10, pady=10)

        def open_third_stage():
            window.destroy()
            self.show_third_stage_window()

        ctk.CTkButton(
            controls_frame,
            text="Text Correction",
            command=open_third_stage,
            width=180,
            height=40,
            font=ctk.CTkFont(size=13),
            fg_color="orange",
            hover_color="darkorange",
        ).pack(side="left", padx=10, pady=10)

        ctk.CTkButton(
            controls_frame,
            text="Close",
            command=window.destroy,
            width=100,
            height=40,
        ).pack(side="left", padx=10, pady=10)

    def process_json_by_parts_worker(self, parent_window, start_button):
        """
        Background worker to process an existing final_output.json part-by-part
        using the second-stage prompt.
        """
        try:
            prompt = self.second_stage_prompt_text.get("1.0", tk.END).strip()
            if not prompt:
                messagebox.showerror("Error", "Please enter a prompt for second-stage processing.")
                return

            # Get chapter name and replace placeholder in prompt
            chapter_name = self.second_stage_chapter_var.get().strip()
            if "{CHAPTER_NAME}" in prompt:
                prompt = prompt.replace("{CHAPTER_NAME}", chapter_name)

            json_path = self.second_stage_json_var.get().strip()
            if not json_path:
                if self.last_final_output_path and os.path.exists(self.last_final_output_path):
                    json_path = self.last_final_output_path
                    self.second_stage_json_var.set(json_path)
                else:
                    messagebox.showerror("Error", "Please select the input JSON file.")
                    return

            if not os.path.exists(json_path):
                messagebox.showerror("Error", f"JSON file not found:\n{json_path}")
                return

            # Use selected model for second stage
            model_name = self.second_stage_model_var.get()

            self.update_status("Starting second-stage part processing...")

            final_path = self.multi_part_post_processor.process_final_json_by_parts(
                json_path=json_path,
                user_prompt=prompt,
                model_name=model_name,
            )

            if not final_path or not os.path.exists(final_path):
                self.update_status("❌ Second-stage part processing failed.")
                messagebox.showerror("Error", "Second-stage part processing failed. Check logs for details.")
                return

            # Load and display final combined JSON
            try:
                with open(final_path, "r", encoding="utf-8") as f:
                    final_data = json.load(f)
            except Exception as e:
                self.update_status(f"Error loading post-processed JSON: {str(e)}")
                messagebox.showerror("Error", f"Failed to load post-processed JSON:\n{str(e)}")
                return

            rows = final_data.get("rows", [])
            total_rows = len(rows)
            self.update_status(f"✓ Second-stage processing completed. Total rows: {total_rows}")
            
            # Store post-processed path for third stage
            self.last_post_processed_path = final_path

            # Convert second-stage JSON rows to CSV rows and show as table
            csv_rows = self.json_rows_to_csv_rows_generic(rows)
            if csv_rows:
                csv_window = ctk.CTkToplevel(self.root)
                base_name = os.path.basename(final_path)
                csv_window.title(f"Second-Stage CSV View - {base_name}")
                csv_window.geometry("1200x700")
                csv_window.minsize(800, 500)

                self.show_csv_table(csv_window, csv_rows)

                info_label = ctk.CTkLabel(
                    csv_window,
                    text=f"Displaying {len(csv_rows) - 1} rows from {base_name}",
                    font=ctk.CTkFont(size=12),
                )
                info_label.pack(pady=5)
            else:
                # Fallback: show raw JSON if conversion failed
                response_text = json.dumps(final_data, ensure_ascii=False, indent=2)
                self.show_response_window(response_text, final_path, False, True)

        except Exception as e:
            self.logger.error(f"Error in second-stage processing: {str(e)}", exc_info=True)
            messagebox.showerror("Error", f"Second-stage processing error:\n{str(e)}")
        finally:
            # Re-enable start button on UI thread
            try:
                self.root.after(
                    0,
                    lambda: start_button.configure(state="normal", text="🚀 Start Part Processing")
                )
            except Exception:
                pass

    def show_third_stage_window(self):
        """Show third stage window for text correction"""
        window = ctk.CTkToplevel(self.root)
        window.title("Third Stage - Text Correction")
        window.geometry("900x800")
        window.minsize(800, 700)

        main_frame = ctk.CTkScrollableFrame(window)
        main_frame.pack(fill="both", expand=True, padx=20, pady=20)

        title = ctk.CTkLabel(
            main_frame,
            text="Third Stage - Text Correction",
            font=ctk.CTkFont(size=22, weight="bold"),
        )
        title.pack(pady=(0, 20))

        # First JSON (from stage 1)
        json1_frame = ctk.CTkFrame(main_frame)
        json1_frame.pack(fill="x", pady=(0, 15))

        ctk.CTkLabel(
            json1_frame,
            text="📄 First JSON (from Stage 1)",
            font=ctk.CTkFont(size=16, weight="bold"),
        ).pack(pady=(15, 10))

        file1_frame = ctk.CTkFrame(json1_frame)
        file1_frame.pack(fill="x", padx=15, pady=5)

        ctk.CTkLabel(
            file1_frame,
            text="JSON File (Stage 1):",
            font=ctk.CTkFont(size=12, weight="bold"),
        ).pack(anchor="w", padx=10, pady=(10, 5))

        self.third_stage_json1_var = ctk.StringVar()
        if self.last_final_output_path and os.path.exists(self.last_final_output_path):
            self.third_stage_json1_var.set(self.last_final_output_path)

        json1_entry = ctk.CTkEntry(file1_frame, textvariable=self.third_stage_json1_var, width=400)
        json1_entry.pack(side="left", fill="x", expand=True, padx=(10, 5), pady=(0, 5))

        def browse_json1():
            filename = filedialog.askopenfilename(
                title="Select Stage 1 JSON file",
                filetypes=[("JSON files", "*.json"), ("All files", "*.*")]
            )
            if filename:
                self.third_stage_json1_var.set(filename)

        ctk.CTkButton(
            file1_frame,
            text="Browse",
            command=browse_json1,
            width=80,
        ).pack(side="right", padx=(5, 10), pady=(0, 5))

        # Second JSON (from stage 2)
        json2_frame = ctk.CTkFrame(main_frame)
        json2_frame.pack(fill="x", pady=(0, 15))

        ctk.CTkLabel(
            json2_frame,
            text="📄 Second JSON (from Stage 2)",
            font=ctk.CTkFont(size=16, weight="bold"),
        ).pack(pady=(15, 10))

        file2_frame = ctk.CTkFrame(json2_frame)
        file2_frame.pack(fill="x", padx=15, pady=5)

        ctk.CTkLabel(
            file2_frame,
            text="JSON File (Stage 2):",
            font=ctk.CTkFont(size=12, weight="bold"),
        ).pack(anchor="w", padx=10, pady=(10, 5))

        self.third_stage_json2_var = ctk.StringVar()
        if self.last_post_processed_path and os.path.exists(self.last_post_processed_path):
            self.third_stage_json2_var.set(self.last_post_processed_path)

        json2_entry = ctk.CTkEntry(file2_frame, textvariable=self.third_stage_json2_var, width=400)
        json2_entry.pack(side="left", fill="x", expand=True, padx=(10, 5), pady=(0, 5))

        def browse_json2():
            filename = filedialog.askopenfilename(
                title="Select Stage 2 JSON file",
                filetypes=[("JSON files", "*.json"), ("All files", "*.*")]
            )
            if filename:
                self.third_stage_json2_var.set(filename)

        ctk.CTkButton(
            file2_frame,
            text="Browse",
            command=browse_json2,
            width=80,
        ).pack(side="right", padx=(5, 10), pady=(0, 5))

        # Prompt section
        prompt_frame = ctk.CTkFrame(main_frame)
        prompt_frame.pack(fill="x", pady=(0, 15))

        ctk.CTkLabel(
            prompt_frame,
            text="💬 Correction Prompt",
            font=ctk.CTkFont(size=16, weight="bold"),
        ).pack(pady=(15, 10))

        ctk.CTkLabel(
            prompt_frame,
            text="Enter the prompt for text correction. Both JSON files will be sent to the model.",
            font=ctk.CTkFont(size=11),
            text_color="gray",
        ).pack(anchor="w", padx=10, pady=(0, 10))

        self.third_stage_prompt_text = ctk.CTkTextbox(prompt_frame, height=140)
        self.third_stage_prompt_text.pack(fill="x", padx=10, pady=(0, 10))

        # Chapter name section (right after prompt)
        chapter_frame = ctk.CTkFrame(prompt_frame)
        chapter_frame.pack(fill="x", padx=10, pady=(0, 10))

        ctk.CTkLabel(
            chapter_frame,
            text="Chapter Name:",
            font=ctk.CTkFont(size=12, weight="bold"),
        ).pack(anchor="w", padx=10, pady=(10, 5))

        self.third_stage_chapter_var = ctk.StringVar(value="درمان با UV")
        chapter_entry = ctk.CTkEntry(
            chapter_frame,
            textvariable=self.third_stage_chapter_var,
            width=400
        )
        chapter_entry.pack(anchor="w", padx=10, pady=(0, 5))

        ctk.CTkLabel(
            chapter_frame,
            text="This name will automatically replace {CHAPTER_NAME} in your prompt.",
            font=ctk.CTkFont(size=10),
            text_color="gray",
        ).pack(anchor="w", padx=10, pady=(0, 10))

        # Start PointId section
        ctk.CTkLabel(
            chapter_frame,
            text="Start PointId (e.g. 1050030000):",
            font=ctk.CTkFont(size=12, weight="bold"),
        ).pack(anchor="w", padx=10, pady=(5, 5))

        self.third_stage_pointid_var = ctk.StringVar(value="")
        pointid_entry = ctk.CTkEntry(
            chapter_frame,
            textvariable=self.third_stage_pointid_var,
            width=400
        )
        pointid_entry.pack(anchor="w", padx=10, pady=(0, 5))

        ctk.CTkLabel(
            chapter_frame,
            text="If empty, a default starting PointId will be used.",
            font=ctk.CTkFont(size=10),
            text_color="gray",
        ).pack(anchor="w", padx=10, pady=(0, 5))

        # Model selection
        model_frame = ctk.CTkFrame(main_frame)
        model_frame.pack(fill="x", pady=(0, 15))

        ctk.CTkLabel(
            model_frame,
            text="🤖 Model Selection",
            font=ctk.CTkFont(size=16, weight="bold"),
        ).pack(pady=(15, 10))

        inner_model_frame = ctk.CTkFrame(model_frame)
        inner_model_frame.pack(fill="x", padx=15, pady=5)

        ctk.CTkLabel(
            inner_model_frame,
            text="Select Model:",
            font=ctk.CTkFont(size=12, weight="bold"),
        ).pack(anchor="w", padx=10, pady=(10, 5))

        self.third_stage_model_var = ctk.StringVar(value=self.model_var.get())
        self.third_stage_model_combo = ctk.CTkComboBox(
            inner_model_frame,
            values=APIConfig.TEXT_MODELS,
            variable=self.third_stage_model_var,
            width=400,
        )
        self.third_stage_model_combo.pack(anchor="w", padx=10, pady=(0, 10))

        # Control buttons
        controls_frame = ctk.CTkFrame(main_frame)
        controls_frame.pack(fill="x", pady=(10, 10))

        def start_third_stage():
            start_button.configure(state="disabled", text="Processing...")
            threading.Thread(
                target=self.process_third_stage_worker,
                args=(window, start_button),
                daemon=True
            ).start()

        start_button = ctk.CTkButton(
            controls_frame,
            text="Start Text Correction",
            command=start_third_stage,
            width=220,
            height=40,
            font=ctk.CTkFont(size=14, weight="bold"),
        )
        start_button.pack(side="left", padx=10, pady=10)

        # View/Edit Stage 3 JSON button
        ctk.CTkButton(
            controls_frame,
            text="View/Edit Stage 3 JSON",
            command=lambda: self.open_json_editor(self.last_corrected_path),
            width=200,
            height=40,
            font=ctk.CTkFont(size=13),
        ).pack(side="left", padx=10, pady=10)

        ctk.CTkButton(
            controls_frame,
            text="Close",
            command=window.destroy,
            width=100,
            height=40,
        ).pack(side="left", padx=10, pady=10)

    def process_third_stage_worker(self, parent_window, start_button):
        """Background worker for third stage text correction"""
        try:
            json1_path = self.third_stage_json1_var.get().strip()
            json2_path = self.third_stage_json2_var.get().strip()
            prompt = self.third_stage_prompt_text.get("1.0", tk.END).strip()
            model_name = self.third_stage_model_var.get()

            # Get chapter name and replace placeholder in prompt
            chapter_name = self.third_stage_chapter_var.get().strip()
            if "{CHAPTER_NAME}" in prompt:
                prompt = prompt.replace("{CHAPTER_NAME}", chapter_name)

            # Get starting PointId from user (optional)
            start_pointid_str = self.third_stage_pointid_var.get().strip()
            if not start_pointid_str:
                # Default for single-chapter mode (can be adjusted)
                start_pointid_str = "1050030000"

            if not (start_pointid_str.isdigit() and len(start_pointid_str) == 10):
                messagebox.showerror("Error", "Start PointId must be a 10-digit number, e.g. 1050030000.")
                return

            try:
                book_id = int(start_pointid_str[0:3])
                chapter_id_num = int(start_pointid_str[3:6])
                start_index = int(start_pointid_str[6:10])
            except ValueError:
                messagebox.showerror("Error", "Start PointId format is invalid.")
                return

            # Validate inputs
            if not json1_path or not os.path.exists(json1_path):
                messagebox.showerror("Error", "Please select a valid Stage 1 JSON file.")
                return

            if not json2_path or not os.path.exists(json2_path):
                messagebox.showerror("Error", "Please select a valid Stage 2 JSON file.")
                return

            if not prompt:
                messagebox.showerror("Error", "Please enter a correction prompt.")
                return

            self.update_status("Loading JSON files...")

            # Load JSON files
            try:
                with open(json1_path, "r", encoding="utf-8") as f:
                    json1_data = json.load(f)
                with open(json2_path, "r", encoding="utf-8") as f:
                    json2_data = json.load(f)
            except Exception as e:
                messagebox.showerror("Error", f"Failed to load JSON files:\n{str(e)}")
                return

            self.update_status("Preparing prompt with JSON data...")

            # Prepare combined prompt
            json1_str = json.dumps(json1_data, ensure_ascii=False, indent=2)
            json2_str = json.dumps(json2_data, ensure_ascii=False, indent=2)

            full_prompt = f"""{prompt}

First JSON (from Stage 1):
{json1_str}

Second JSON (from Stage 2):
{json2_str}

Please process both JSON files according to the prompt above and return the result as JSON."""

            self.update_status(f"Processing with {model_name}...")

            # Process with model
            response = self.api_client.process_text(
                text=full_prompt,
                model_name=model_name,
                temperature=0.7,
                max_tokens=APIConfig.DEFAULT_MAX_TOKENS,
            )

            if not response:
                self.update_status("❌ Text correction failed.")
                messagebox.showerror("Error", "Text correction failed. Check logs for details.")
                return

            self.update_status("Parsing response...")

            # Try to extract JSON from response
            try:
                # Clean response
                cleaned = response.strip()
                if cleaned.startswith('```'):
                    end_idx = cleaned.find('```', 3)
                    if end_idx > 0:
                        cleaned = cleaned[3:end_idx].strip()
                        if cleaned.startswith('json'):
                            cleaned = cleaned[4:].strip()

                result_data = json.loads(cleaned)
            except json.JSONDecodeError:
                # If not JSON, wrap in a structure
                result_data = {
                    "response": response,
                    "processed_at": datetime.now().isoformat(),
                    "model": model_name,
                }

            # Save raw result (for debugging, not shown to user)
            base_name1 = os.path.splitext(os.path.basename(json1_path))[0]
            base_name2 = os.path.splitext(os.path.basename(json2_path))[0]
            raw_output_filename = f"{base_name1}_{base_name2}_corrected_raw.json"
            raw_output_path = os.path.join(os.path.dirname(json1_path), raw_output_filename)

            try:
                with open(raw_output_path, "w", encoding="utf-8") as f:
                    json.dump(result_data, f, ensure_ascii=False, indent=2)
            except Exception as e:
                self.logger.warning(f"Failed to save raw corrected JSON: {str(e)}")

            # Post-process result_data into clean hierarchical structure with PointId
            self.update_status("Post-processing corrected JSON into final structured output...")

            def clean_label(text: str) -> str:
                if not isinstance(text, str):
                    return ""
                prefixes = ["فصل:", "زیرفصل:", "مبحث:", "عنوان:", "زیرعنوان:"]
                for pref in prefixes:
                    if text.startswith(pref):
                        return text[len(pref):].strip()
                return text.strip()

            def clean_point_text(text: str) -> str:
                if not isinstance(text, str):
                    return ""
                return text.lstrip("•").strip()

            # Determine inner structured JSON
            inner = None
            if isinstance(result_data, dict) and "chapter" in result_data and "content" in result_data:
                inner = result_data
            elif isinstance(result_data, dict) and isinstance(result_data.get("response"), str):
                inner_text = result_data["response"].strip()

                # Remove markdown code fences if present
                # Handle cases like: ```json\n{...}\n``` or ```\n{...}\n```
                if "```" in inner_text:
                    # Find all code fence positions
                    fence_positions = []
                    idx = 0
                    while True:
                        pos = inner_text.find("```", idx)
                        if pos == -1:
                            break
                        fence_positions.append(pos)
                        idx = pos + 3
                    
                    # If we have at least 2 fences, extract content between first and last
                    if len(fence_positions) >= 2:
                        start_fence = fence_positions[0] + 3  # After first ```
                        end_fence = fence_positions[-1]  # Before last ```
                        inner_text = inner_text[start_fence:end_fence].strip()
                
                # Remove leading "json" identifier if present
                if inner_text.lower().startswith("json"):
                    inner_text = inner_text[4:].strip()

                # Try multiple strategies to parse JSON
                inner = None
                
                # Strategy 1: Direct JSON parse
                try:
                    inner = json.loads(inner_text)
                except Exception as e1:
                    # Strategy 2: Extract first {...} block (most reliable)
                    try:
                        start = inner_text.find("{")
                        end = inner_text.rfind("}")
                        if start != -1 and end != -1 and end > start:
                            candidate = inner_text[start:end + 1]
                            inner = json.loads(candidate)
                        else:
                            self.logger.warning(f"Could not find JSON boundaries. Start: {start}, End: {end}")
                    except Exception as e2:
                        self.logger.error(f"Failed to parse JSON from response. Direct parse error: {str(e1)}, Extract parse error: {str(e2)}")
                        inner = None

            if not isinstance(inner, dict) or "content" not in inner:
                # Fallback: show raw JSON if we cannot interpret structure
                error_msg = "Could not parse hierarchical structure from model response."
                if inner is None:
                    error_msg += "\n\nReason: Failed to extract JSON from response string."
                elif not isinstance(inner, dict):
                    error_msg += f"\n\nReason: Extracted data is not a dictionary (type: {type(inner).__name__})."
                elif "content" not in inner:
                    error_msg += f"\n\nReason: Missing 'content' field in extracted JSON. Available keys: {list(inner.keys())[:5]}"
                
                self.update_status(f"Warning: {error_msg}")
                self.logger.warning(error_msg)
                
                output_filename = f"{base_name1}_{base_name2}_corrected.json"
                output_path = os.path.join(os.path.dirname(json1_path), output_filename)
                try:
                    with open(output_path, "w", encoding="utf-8") as f:
                        json.dump(result_data, f, ensure_ascii=False, indent=2)
                except Exception as e:
                    self.logger.error(f"Failed to save fallback corrected JSON: {str(e)}")

                self.last_corrected_path = output_path
                response_text = json.dumps(result_data, ensure_ascii=False, indent=2)
                self.show_response_window(response_text, output_path, False, True)
                messagebox.showwarning("Warning", f"{error_msg}\n\nRaw JSON saved to:\n{output_path}")
                return

            chapter_internal = inner.get("chapter", chapter_name)
            content = inner.get("content", [])

            flat_rows = []

            def walk(node, lvl1=None, lvl2=None, lvl3=None, lvl4=None):
                l1 = lvl1
                l2 = lvl2
                l3 = lvl3
                l4 = lvl4

                if "level_1" in node:
                    l1 = clean_label(node["level_1"])
                if "level_2" in node:
                    l2 = clean_label(node["level_2"])
                if "level_3" in node:
                    l3 = clean_label(node["level_3"])
                if "level_4" in node:
                    l4 = clean_label(node["level_4"])

                if "level_5" in node:
                    subsub = clean_label(node["level_5"])
                    for p in node.get("points", []):
                        flat_rows.append({
                            "chapter": chapter_internal,
                            "subchapter": l2 or "",
                            "topic": l3 or "",
                            "subtopic": l4 or "",
                            "subsubtopic": subsub,
                            "points": clean_point_text(p),
                        })

                for child in node.get("children", []):
                    if isinstance(child, dict):
                        walk(child, l1, l2, l3, l4)

            for item in content:
                if isinstance(item, dict):
                    walk(item, None, None, None, None)

            if not flat_rows:
                self.update_status("Warning: No points extracted from corrected JSON. Showing raw JSON.")
                output_filename = f"{base_name1}_{base_name2}_corrected.json"
                output_path = os.path.join(os.path.dirname(json1_path), output_filename)
                try:
                    with open(output_path, "w", encoding="utf-8") as f:
                        json.dump(result_data, f, ensure_ascii=False, indent=2)
                except Exception as e:
                    self.logger.error(f"Failed to save fallback corrected JSON: {str(e)}")

                self.last_corrected_path = output_path
                response_text = json.dumps(result_data, ensure_ascii=False, indent=2)
                self.show_response_window(response_text, output_path, False, True)
                messagebox.showinfo("Warning", f"No structured points found. Raw JSON saved to:\n{output_path}")
                return

            # Generate PointId for each row
            current_index = start_index
            for row in flat_rows:
                point_id = f"{book_id:03d}{chapter_id_num:03d}{current_index:04d}"
                row["PointId"] = point_id
                current_index += 1

            # Build final clean JSON
            clean_output = {
                "metadata": {
                    "chapter": chapter_internal,
                    "book_id": book_id,
                    "chapter_id": chapter_id_num,
                    "start_point_index": start_index,
                    "total_points": len(flat_rows),
                    "processed_at": datetime.now().isoformat(),
                    "source_stage1": os.path.basename(json1_path),
                    "source_stage2": os.path.basename(json2_path),
                    "model": model_name,
                },
                "points": flat_rows,
            }

            # Save clean output
            output_filename = f"{base_name1}_{base_name2}_corrected.json"
            output_path = os.path.join(os.path.dirname(json1_path), output_filename)

            with open(output_path, "w", encoding="utf-8") as f:
                json.dump(clean_output, f, ensure_ascii=False, indent=2)

            self.update_status(f"✓ Text correction completed. Final structured JSON saved to: {output_path}")

            # Store corrected JSON path for later editing
            self.last_corrected_path = output_path

            # Display clean result
            response_text = json.dumps(clean_output, ensure_ascii=False, indent=2)
            self.show_response_window(response_text, output_path, False, True)

            messagebox.showinfo("Success", f"Text correction completed!\n\nFinal JSON saved to:\n{output_path}")

        except Exception as e:
            self.logger.error(f"Error in third-stage processing: {str(e)}", exc_info=True)
            messagebox.showerror("Error", f"Text correction error:\n{str(e)}")
        finally:
            try:
                self.root.after(
                    0,
                    lambda: start_button.configure(state="normal", text="🚀 Start Text Correction")
                )
            except Exception:
                pass
    
    def is_json_format(self, text: str) -> bool:
        """
        Check if the response text is in JSON format
        
        Args:
            text: Response text to check
            
        Returns:
            True if text appears to be JSON format
        """
        if not text or not text.strip():
            return False
        
        # Try to parse as JSON
        try:
            # Remove markdown code blocks if present
            cleaned_text = text.strip()
            if cleaned_text.startswith('```'):
                # Extract content from code block
                lines = cleaned_text.split('\n')
                if len(lines) > 1:
                    # Skip first line (```json or ```)
                    cleaned_text = '\n'.join(lines[1:])
                    # Remove last line if it's ```
                    if cleaned_text.endswith('```'):
                        cleaned_text = cleaned_text[:-3].strip()
            
            # Try to parse as JSON
            json.loads(cleaned_text)
            return True
        except (json.JSONDecodeError, ValueError):
            # Check if it looks like JSON (starts with [ or {)
            cleaned_text = text.strip()
            if cleaned_text.startswith('[') or cleaned_text.startswith('{'):
                # Try to extract JSON from text
                try:
                    # Look for JSON array or object
                    if cleaned_text.startswith('['):
                        # Find matching closing bracket
                        bracket_count = 0
                        for i, char in enumerate(cleaned_text):
                            if char == '[':
                                bracket_count += 1
                            elif char == ']':
                                bracket_count -= 1
                                if bracket_count == 0:
                                    json_text = cleaned_text[:i+1]
                                    json.loads(json_text)
                                    return True
                    elif cleaned_text.startswith('{'):
                        # Find matching closing brace
                        brace_count = 0
                        for i, char in enumerate(cleaned_text):
                            if char == '{':
                                brace_count += 1
                            elif char == '}':
                                brace_count -= 1
                                if brace_count == 0:
                                    json_text = cleaned_text[:i+1]
                                    json.loads(json_text)
                                    return True
                except (json.JSONDecodeError, ValueError):
                    pass
        
        return False
    
    def parse_json_to_table_data(self, text: str) -> Optional[List[List[str]]]:
        """
        Parse JSON text into rows and columns for table display
        
        Args:
            text: JSON text to parse
            
        Returns:
            List of rows (each row is a list of cells) or None if parsing fails
        """
        try:
            # Remove markdown code blocks if present
            cleaned_text = text.strip()
            if cleaned_text.startswith('```'):
                lines = cleaned_text.split('\n')
                if len(lines) > 1:
                    cleaned_text = '\n'.join(lines[1:])
                    if cleaned_text.endswith('```'):
                        cleaned_text = cleaned_text[:-3].strip()
            
            # Parse JSON
            json_data = json.loads(cleaned_text)
            
            # Handle list of objects (most common case)
            if isinstance(json_data, list) and len(json_data) > 0:
                # Get all unique keys from all objects
                all_keys = set()
                for item in json_data:
                    if isinstance(item, dict):
                        all_keys.update(item.keys())
                
                # Sort keys for consistent column order
                headers = sorted(list(all_keys))
                
                # Create rows: first row is headers, then data rows
                rows = [headers]  # Header row
                
                for item in json_data:
                    if isinstance(item, dict):
                        row = [str(item.get(key, "")) for key in headers]
                        rows.append(row)
                    else:
                        # If item is not a dict, convert to string
                        rows.append([str(item)])
                
                return rows if len(rows) > 1 else None  # At least header + one data row
            
            # Handle single object
            elif isinstance(json_data, dict):
                headers = list(json_data.keys())
                rows = [headers]  # Header row
                row = [str(json_data.get(key, "")) for key in headers]
                rows.append(row)
                return rows
            
            # Handle other types (convert to string)
            else:
                return [[str(json_data)]]
                
        except (json.JSONDecodeError, ValueError) as e:
            self.logger.error(f"Error parsing JSON: {str(e)}")
            return None
        except Exception as e:
            self.logger.error(f"Unexpected error parsing JSON: {str(e)}")
            return None
    
    def process_pdf(self):
        """Process PDF with selected prompt and model"""
        def worker():
            try:
                # Disable process button
                self.process_btn.configure(state="disabled", text="Processing...")
                
                # Validate inputs
                if not self.pdf_path or not os.path.exists(self.pdf_path):
                    self.update_status("Error: Please select a valid PDF file")
                    messagebox.showerror("Error", "Please select a valid PDF file")
                    return
                
                prompt = self.get_selected_prompt()
                if not prompt:
                    self.update_status("Error: Please select or enter a prompt")
                    messagebox.showerror("Error", "Please select or enter a prompt")
                    return
                
                # Validate API keys are loaded
                if not self.api_key_manager.api_keys:
                    if not self.api_key_file_var.get():
                        self.update_status("Error: Please load API keys from CSV file")
                        messagebox.showerror("Error", "Please load API keys from CSV file")
                        return
                    else:
                        if not self.api_key_manager.load_from_csv(self.api_key_file_var.get()):
                            self.update_status("Error: Failed to load API keys")
                            messagebox.showerror("Error", "Failed to load API keys from CSV file")
                            return
                
                # Get selected model
                model_name = self.model_var.get()
                
                # Get current API key info (for logging)
                current_key_info = self.api_key_manager.get_current_key_info()
                if current_key_info:
                    self.update_status(f"Using API key: {current_key_info['account']} ({len(self.api_key_manager.api_keys)} keys available)")
                
                self.update_status(f"Processing PDF with {model_name}...")
                self.update_status(f"Prompt length: {len(prompt)} characters")
                self.update_status(f"PDF file: {os.path.basename(self.pdf_path)}")
                self.logger.info(f"=== Starting PDF Processing ===")
                self.logger.info(f"Model: {model_name}")
                self.logger.info(f"PDF: {self.pdf_path}")
                self.logger.info(f"Full prompt being sent ({len(prompt)} chars): {prompt}")
                
                # Process PDF using multi-part processor
                # This handles large outputs by splitting into multiple parts
                self.update_status("Processing PDF with multi-part system...")
                self.update_status("Large outputs will be split into multiple parts")
                self.update_status("Only the final combined JSON will be displayed")
                
                def progress_callback(message: str):
                    """Callback for progress updates during multi-part processing"""
                    self.update_status(message)
                
                # Use multi-part processor
                final_output_path = self.multi_part_processor.process_multi_part(
                    pdf_path=self.pdf_path,
                    base_prompt=prompt,  # User's original prompt (no modifications)
                    model_name=model_name,
                    temperature=0.7,
                    resume=True,  # Enable resume capability
                    progress_callback=progress_callback
                )
                
                if final_output_path and os.path.exists(final_output_path):
                    # Load final output JSON
                    try:
                        with open(final_output_path, 'r', encoding='utf-8') as f:
                            final_data = json.load(f)
                        
                        # Extract metadata and rows
                        metadata = final_data.get('metadata', {})
                        rows = final_data.get('rows', [])
                        
                        total_parts = metadata.get('total_parts', 0)
                        total_rows = metadata.get('total_rows', len(rows))
                        
                        self.update_status("✓ Multi-part processing completed successfully!")
                        self.update_status(f"Total parts processed: {total_parts}")
                        self.update_status(f"Total rows extracted: {total_rows}")
                        self.update_status(f"Final output saved to: {final_output_path}")
                        
                        # Store final output path for CSV viewing
                        self.last_final_output_path = final_output_path
                        
                        # Display final output (not individual parts)
                        if rows:
                            self.update_status(f"\n📄 Displaying final combined output ({total_rows} rows)...")
                            # Show JSON table
                            is_json = True
                            response_text = json.dumps(final_data, ensure_ascii=False, indent=2)
                            self.show_response_window(response_text, final_output_path, False, is_json)
                        else:
                            self.update_status("⚠ Warning: Final output contains no rows!")
                            messagebox.showwarning("No Data", "The final output file contains no rows. Check the logs for details.")
                            return
                        
                    except json.JSONDecodeError as e:
                        self.update_status(f"Error: Failed to parse final output JSON: {str(e)}")
                        self.logger.error(f"Failed to parse final output: {str(e)}")
                        messagebox.showerror("Error", f"Failed to parse final output JSON: {str(e)}")
                        return
                    except Exception as e:
                        self.update_status(f"Error loading final output: {str(e)}")
                        self.logger.error(f"Error loading final output: {str(e)}", exc_info=True)
                        messagebox.showerror("Error", f"Failed to load final output: {str(e)}")
                        return
                else:
                    self.update_status("❌ Multi-part processing failed or was incomplete")
                    self.logger.error("Multi-part processing returned no final output file")
                    messagebox.showerror("Error", "Multi-part processing failed. Check the logs for details.")
                    return
                
            except Exception as e:
                error_msg = f"Error processing PDF: {str(e)}"
                self.update_status(error_msg)
                self.logger.error(error_msg, exc_info=True)
                messagebox.showerror("Error", error_msg)
            finally:
                # Re-enable process button
                self.process_btn.configure(state="normal", text="🚀 Process PDF with AI")
        
        # Run in separate thread to prevent UI blocking
        threading.Thread(target=worker, daemon=True).start()
    
    def save_response(self, response: str, is_json: bool = False) -> Optional[str]:
        """
        Save response to file (JSON if JSON format, otherwise Word document)
        
        Args:
            response: Response text to save
            is_json: Whether the response is in JSON format
            
        Returns:
            Path to saved file or None if failed
        """
        try:
            # JSON is already saved by API layer, just return the path
            if is_json:
                # The API layer saves JSON automatically, so we don't need to save again
                # But we can return a message that it's already saved
                self.logger.info("JSON response was already saved by API layer")
                return None
            
            # Get output folder
            output_folder = self.output_folder_var.get().strip()
            if not output_folder:
                # Use default: same folder as PDF or current directory
                if self.pdf_path:
                    output_folder = os.path.dirname(self.pdf_path)
                else:
                    output_folder = os.getcwd()
            
            # Create output folder if it doesn't exist
            if not os.path.exists(output_folder):
                os.makedirs(output_folder)
            
            # Generate filename based on PDF name and timestamp
            from datetime import datetime
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            
            if self.pdf_path:
                pdf_name = os.path.splitext(os.path.basename(self.pdf_path))[0]
                filename = f"{pdf_name}_response_{timestamp}.docx"
            else:
                filename = f"response_{timestamp}.docx"
            
            file_path = os.path.join(output_folder, filename)
            
            # Save as Word document (for non-JSON responses)
            # JSON responses are already saved by API layer, so we only save non-JSON responses here
            try:
                from docx import Document
                from docx.shared import Pt, RGBColor
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                
                doc = Document()
                
                # Title
                title = doc.add_heading('PDF Processing Response', 0)
                title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Metadata section
                doc.add_paragraph().add_run('Metadata').bold = True
                metadata_para = doc.add_paragraph()
                metadata_para.add_run(f'Generated: ').bold = True
                metadata_para.add_run(f"{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
                
                if self.pdf_path:
                    metadata_para = doc.add_paragraph()
                    metadata_para.add_run(f'PDF File: ').bold = True
                    metadata_para.add_run(os.path.basename(self.pdf_path))
                
                metadata_para = doc.add_paragraph()
                metadata_para.add_run(f'Model: ').bold = True
                metadata_para.add_run(self.model_var.get())
                
                # Full prompt
                prompt = self.get_selected_prompt()
                if prompt:
                    metadata_para = doc.add_paragraph()
                    metadata_para.add_run(f'Prompt: ').bold = True
                    doc.add_paragraph(prompt)
                
                # Separator
                doc.add_paragraph('=' * 80)
                
                # Response content
                doc.add_paragraph().add_run('Response').bold = True
                
                # Split response into paragraphs and add them
                response_paragraphs = response.split('\n\n')
                for para_text in response_paragraphs:
                    if para_text.strip():
                        # Preserve formatting for lists and special content
                        if para_text.strip().startswith('-') or para_text.strip().startswith('*'):
                            # List item
                            doc.add_paragraph(para_text.strip(), style='List Bullet')
                        elif para_text.strip().startswith(tuple('123456789')):
                            # Numbered list
                            doc.add_paragraph(para_text.strip(), style='List Number')
                        else:
                            # Regular paragraph
                            doc.add_paragraph(para_text.strip())
                
                # Save document
                doc.save(file_path)
                self.logger.info(f"Response saved to Word document: {file_path}")
                return file_path
                
            except ImportError:
                # Fallback to text file if python-docx is not available
                self.logger.warning("python-docx not available, saving as text file")
                txt_file_path = file_path.replace('.docx', '.txt')
                
                # Get prompt again for fallback
                prompt_text = self.get_selected_prompt() or "N/A"
                
                with open(txt_file_path, 'w', encoding='utf-8') as f:
                    f.write(f"PDF Processing Response\n")
                    f.write(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
                    if self.pdf_path:
                        f.write(f"PDF File: {os.path.basename(self.pdf_path)}\n")
                    f.write(f"Model: {self.model_var.get()}\n")
                    f.write(f"Prompt: {prompt_text}\n")
                    f.write(f"{'='*80}\n\n")
                    f.write(response)
                return txt_file_path
            
        except Exception as e:
            self.logger.error(f"Error saving response: {str(e)}")
            return None
    
    def show_response_window(self, response: str, saved_path: Optional[str] = None, 
                            is_truncated: bool = False, is_json: bool = False):
        """Show full response in a new window (with JSON table view if JSON format)"""
        response_window = ctk.CTkToplevel(self.root)
        response_window.title("AI Response - Full Content" + (" [TRUNCATED]" if is_truncated else "") + (" [JSON]" if is_json else ""))
        # Make window larger for better viewing
        response_window.geometry("1200x800" if is_json else "1000x700")
        
        # Title
        title_text = "AI Response - Full Content"
        if is_json:
            title_text += " [JSON Format]"
        if is_truncated:
            title_text += " ⚠️ [TRUNCATED - INCOMPLETE]"
        title_color = "red" if is_truncated else "white"
        title = ctk.CTkLabel(response_window, text=title_text, 
                            font=ctk.CTkFont(size=20, weight="bold"),
                            text_color=title_color)
        title.pack(pady=10)
        
        # Response length info
        response_length = len(response)
        if is_json:
            try:
                json_data = json.loads(response)
                if isinstance(json_data, list):
                    row_count = len(json_data)
                    length_label = ctk.CTkLabel(response_window, 
                                               text=f"Response length: {response_length:,} characters | JSON Format: {row_count} rows", 
                                               font=ctk.CTkFont(size=11), 
                                               text_color="gray")
                else:
                    length_label = ctk.CTkLabel(response_window, 
                                               text=f"Response length: {response_length:,} characters | JSON Format", 
                                               font=ctk.CTkFont(size=11), 
                                               text_color="gray")
            except:
                length_label = ctk.CTkLabel(response_window, 
                                           text=f"Response length: {response_length:,} characters | JSON Format", 
                                           font=ctk.CTkFont(size=11), 
                                           text_color="gray")
        else:
            length_label = ctk.CTkLabel(response_window, 
                                       text=f"Response length: {response_length:,} characters", 
                                       font=ctk.CTkFont(size=11), 
                                       text_color="gray")
        length_label.pack(pady=2)
        
        # Warning if truncated
        if is_truncated:
            warning_label = ctk.CTkLabel(response_window, 
                                       text="⚠️ WARNING: Response was TRUNCATED due to MAX_TOKENS limit!\n"
                                            "The response is INCOMPLETE. Consider using gemini-2.5-pro model.", 
                                       font=ctk.CTkFont(size=12, weight="bold"),
                                       text_color="red",
                                       wraplength=900)
            warning_label.pack(pady=5)
        
        # Saved path info (will be updated if saved later)
        saved_label = None
        if saved_path:
            saved_label = ctk.CTkLabel(response_window, 
                                       text=f"✓ Saved to: {os.path.basename(saved_path)}", 
                                       font=ctk.CTkFont(size=10), text_color="green")
            saved_label.pack(pady=5)
        
        # Display JSON as table or regular text
        if is_json:
            # Parse JSON and display as table
            try:
                json_data = json.loads(response)
                if isinstance(json_data, list) and len(json_data) > 0:
                    # Convert JSON to table format
                    self.show_json_table(response_window, json_data)
                else:
                    # Fallback to text view if not a list
                    self.show_text_view(response_window, response)
            except json.JSONDecodeError:
                # Fallback to text view if parsing fails
                self.show_text_view(response_window, response)
        else:
            # Regular text view
            self.show_text_view(response_window, response)
        
        # Buttons frame
        buttons_frame = ctk.CTkFrame(response_window)
        buttons_frame.pack(pady=10)
        
        # Save button (if not already saved)
        if saved_path:
            ctk.CTkButton(buttons_frame, text="📁 Open Folder", 
                         command=lambda: self.open_folder(os.path.dirname(saved_path)),
                         width=140).pack(side="left", padx=5)
        else:
            ctk.CTkButton(buttons_frame, text="💾 Save As...", 
                         command=lambda: self.save_response_as(response, response_window, is_json),
                         width=140).pack(side="left", padx=5)
        
        # Copy button
        ctk.CTkButton(buttons_frame, text="📋 Copy All", 
                     command=lambda: self.copy_to_clipboard(response, response_window),
                     width=120).pack(side="left", padx=5)
        
        # Close button
        ctk.CTkButton(buttons_frame, text="Close", command=response_window.destroy,
                     width=100).pack(side="left", padx=5)
        
        # Store reference to update saved path later if needed
        response_window._saved_label = saved_label
        response_window._saved_path = saved_path

    def open_json_editor(self, json_path: Optional[str] = None):
        """
        Open a JSON editor window for a given JSON file.
        If json_path is None or invalid, ask user to select a JSON file.
        """
        try:
            if not json_path or not os.path.exists(json_path):
                json_path = filedialog.askopenfilename(
                    title="Select JSON file",
                    filetypes=[("JSON files", "*.json"), ("All files", "*.*")]
                )
                if not json_path or not os.path.exists(json_path):
                    messagebox.showwarning("File Not Found", "No valid JSON file selected.")
                    return

            # Load JSON content
            try:
                with open(json_path, "r", encoding="utf-8") as f:
                    data = json.load(f)
                json_text = json.dumps(data, ensure_ascii=False, indent=2)
            except Exception as e:
                messagebox.showerror("Error", f"Failed to load JSON file:\n{str(e)}")
                return

            # Create editor window
            editor = ctk.CTkToplevel(self.root)
            editor.title(f"JSON Editor - {os.path.basename(json_path)}")
            editor.geometry("1000x800")

            # Path label
            path_label = ctk.CTkLabel(
                editor,
                text=f"File: {json_path}",
                font=ctk.CTkFont(size=11),
                text_color="gray",
            )
            path_label.pack(pady=(10, 5))

            # Textbox
            text_frame = ctk.CTkFrame(editor)
            text_frame.pack(fill="both", expand=True, padx=20, pady=10)

            text_widget = ctk.CTkTextbox(
                text_frame,
                height=600,
                wrap="word",
                font=self.farsi_text_font,
            )
            text_widget.pack(fill="both", expand=True)
            text_widget.insert("1.0", json_text)

            # Buttons
            button_frame = ctk.CTkFrame(editor)
            button_frame.pack(pady=10)

            def save_json():
                content = text_widget.get("1.0", tk.END).strip()
                if not content:
                    messagebox.showerror("Error", "JSON content is empty.")
                    return
                try:
                    parsed = json.loads(content)
                except json.JSONDecodeError as e:
                    messagebox.showerror("JSON Error", f"Invalid JSON:\n{str(e)}")
                    return
                try:
                    with open(json_path, "w", encoding="utf-8") as f:
                        json.dump(parsed, f, ensure_ascii=False, indent=2)
                    messagebox.showinfo("Saved", f"JSON saved to:\n{json_path}")
                except Exception as e:
                    messagebox.showerror("Error", f"Failed to save JSON:\n{str(e)}")

            def save_as_json():
                content = text_widget.get("1.0", tk.END).strip()
                if not content:
                    messagebox.showerror("Error", "JSON content is empty.")
                    return
                try:
                    parsed = json.loads(content)
                except json.JSONDecodeError as e:
                    messagebox.showerror("JSON Error", f"Invalid JSON:\n{str(e)}")
                    return
                filename = filedialog.asksaveasfilename(
                    title="Save JSON As",
                    defaultextension=".json",
                    filetypes=[("JSON files", "*.json"), ("All files", "*.*")]
                )
                if not filename:
                    return
                try:
                    with open(filename, "w", encoding="utf-8") as f:
                        json.dump(parsed, f, ensure_ascii=False, indent=2)
                    messagebox.showinfo("Saved", f"JSON saved to:\n{filename}")
                except Exception as e:
                    messagebox.showerror("Error", f"Failed to save JSON:\n{str(e)}")

            ctk.CTkButton(
                button_frame,
                text="Save",
                command=save_json,
                width=120,
            ).pack(side="left", padx=5)

            ctk.CTkButton(
                button_frame,
                text="Save As...",
                command=save_as_json,
                width=120,
            ).pack(side="left", padx=5)

            ctk.CTkButton(
                button_frame,
                text="Close",
                command=editor.destroy,
                width=100,
            ).pack(side="left", padx=5)

        except Exception as e:
            self.logger.error(f"Error opening JSON editor: {str(e)}", exc_info=True)
            messagebox.showerror("Error", f"Could not open JSON editor:\n{str(e)}")
    
    def show_json_table(self, parent_window, json_rows: List[Dict[str, Any]]):
        """Display JSON data as a table using Treeview"""
        # Convert JSON to table format
        if not json_rows:
            return
        
        # Get all unique keys from all rows (columns)
        all_keys = set()
        for row in json_rows:
            if isinstance(row, dict):
                all_keys.update(row.keys())
        
        headers = list(all_keys)
        if not headers:
            self.show_text_view(parent_window, json.dumps(json_rows, indent=2, ensure_ascii=False))
            return
        
        # Limit columns for display
        max_columns = 20
        headers = headers[:max_columns]
        
        # Create frame for table
        table_frame = ctk.CTkFrame(parent_window)
        table_frame.pack(fill="both", expand=True, padx=20, pady=10)
        
        # Create scrollable frame
        scrollable_frame = ctk.CTkScrollableFrame(table_frame)
        scrollable_frame.pack(fill="both", expand=True)
        
        # Create Treeview for table display
        tree_frame = tk.Frame(scrollable_frame)
        tree_frame.pack(fill="both", expand=True)
        
        # Create Treeview
        tree = ttk.Treeview(tree_frame, columns=[f"col{i}" for i in range(len(headers))], 
                           show="headings", height=20)
        
        # Configure columns
        for i, header in enumerate(headers):
            tree.heading(f"col{i}", text=header[:30])  # Limit header length
            tree.column(f"col{i}", width=150, anchor="w")
        
        # Insert data rows
        for row in json_rows:
            if isinstance(row, dict):
                display_row = [str(row.get(h, ""))[:100] for h in headers]  # Limit cell length
                tree.insert("", "end", values=display_row)
        
        # Add scrollbars
        vsb = ttk.Scrollbar(tree_frame, orient="vertical", command=tree.yview)
        hsb = ttk.Scrollbar(tree_frame, orient="horizontal", command=tree.xview)
        tree.configure(yscrollcommand=vsb.set, xscrollcommand=hsb.set)
        
        # Pack tree and scrollbars
        tree.grid(row=0, column=0, sticky="nsew")
        vsb.grid(row=0, column=1, sticky="ns")
        hsb.grid(row=1, column=0, sticky="ew")
        tree_frame.grid_rowconfigure(0, weight=1)
        tree_frame.grid_columnconfigure(0, weight=1)
        
        # Info label
        info_text = f"Displaying {len(json_rows)} rows, {len(headers)} columns"
        if len(all_keys) > max_columns:
            info_text += f" (showing first {max_columns} columns)"
        info_label = ctk.CTkLabel(table_frame, text=info_text, 
                                 font=ctk.CTkFont(size=10), text_color="gray")
        info_label.pack(pady=5)
    
    def show_csv_table(self, parent_window, csv_rows: List[List[str]]):
        """Display CSV data as a table using Treeview"""
        # Create frame for table
        table_frame = ctk.CTkFrame(parent_window)
        table_frame.pack(fill="both", expand=True, padx=20, pady=10)
        
        # Create scrollable frame
        scrollable_frame = ctk.CTkScrollableFrame(table_frame)
        scrollable_frame.pack(fill="both", expand=True)
        
        # Create Treeview for table display
        # We'll use a regular tkinter Treeview inside CTkFrame
        tree_frame = tk.Frame(scrollable_frame)
        tree_frame.pack(fill="both", expand=True)
        
        # Determine number of columns
        if csv_rows:
            num_columns = len(csv_rows[0])
            # Limit to reasonable number of columns for display
            max_columns = 20
            num_columns = min(num_columns, max_columns)
        else:
            num_columns = 1
        
        # Create Treeview
        tree = ttk.Treeview(tree_frame, columns=[f"col{i}" for i in range(num_columns)], 
                           show="headings", height=20)
        
        # Configure columns
        for i in range(num_columns):
            tree.heading(f"col{i}", text=f"Column {i+1}")
            tree.column(f"col{i}", width=150, anchor="w")
        
        # If first row looks like headers, use them
        if csv_rows and len(csv_rows) > 0:
            first_row = csv_rows[0]
            # Check if first row looks like headers (non-numeric, descriptive)
            is_header = any(cell and not cell.replace('.', '').replace('-', '').isdigit() 
                          for cell in first_row[:num_columns])
            
            if is_header and len(csv_rows) > 1:
                # Use first row as headers
                for i, header in enumerate(first_row[:num_columns]):
                    tree.heading(f"col{i}", text=header[:30])  # Limit header length
                    tree.column(f"col{i}", width=150, anchor="w")
                # Start data from second row
                data_start = 1
            else:
                data_start = 0
            
            # Insert data rows
            for row_idx, row in enumerate(csv_rows[data_start:], start=data_start):
                # Truncate row to num_columns
                display_row = row[:num_columns]
                # Pad row if needed
                while len(display_row) < num_columns:
                    display_row.append("")
                tree.insert("", "end", values=display_row)
        
        # Add scrollbars
        vsb = ttk.Scrollbar(tree_frame, orient="vertical", command=tree.yview)
        hsb = ttk.Scrollbar(tree_frame, orient="horizontal", command=tree.xview)
        tree.configure(yscrollcommand=vsb.set, xscrollcommand=hsb.set)
        
        # Pack tree and scrollbars
        tree.grid(row=0, column=0, sticky="nsew")
        vsb.grid(row=0, column=1, sticky="ns")
        hsb.grid(row=1, column=0, sticky="ew")
        tree_frame.grid_rowconfigure(0, weight=1)
        tree_frame.grid_columnconfigure(0, weight=1)
        
        # Info label
        info_text = f"Displaying {len(csv_rows)} rows, {num_columns} columns"
        if csv_rows and len(csv_rows[0]) > num_columns:
            info_text += f" (showing first {num_columns} columns)"
        info_label = ctk.CTkLabel(table_frame, text=info_text, 
                                 font=ctk.CTkFont(size=10), text_color="gray")
        info_label.pack(pady=5)
    
    def show_text_view(self, parent_window, response: str):
        """Display response as regular text"""
        # Response text - Full content with scrollbar
        # Use a frame to contain the textbox and scrollbar
        text_frame = ctk.CTkFrame(parent_window)
        text_frame.pack(fill="both", expand=True, padx=20, pady=10)
        
        response_text = ctk.CTkTextbox(
            text_frame,
            height=500,
            wrap="word",  # Wrap text at word boundaries
            font=self.farsi_text_font,
        )
        response_text.pack(fill="both", expand=True)
        
        # Insert full response
        response_text.insert("1.0", response)
        response_text.configure(state="normal")  # Allow copying and scrolling
        
        # Scroll to top
        response_text.see("1.0")
    
    def copy_to_clipboard(self, text: str, parent_window):
        """Copy text to clipboard"""
        try:
            parent_window.clipboard_clear()
            parent_window.clipboard_append(text)
            messagebox.showinfo("Copied", "Response copied to clipboard!")
        except Exception as e:
            self.logger.error(f"Error copying to clipboard: {str(e)}")
            messagebox.showerror("Error", f"Could not copy to clipboard: {str(e)}")
    
    def save_response_as(self, response: str, parent_window, is_json: bool = False):
        """Save response to a custom location"""
        if is_json:
            filetypes = [("JSON files", "*.json"), ("Text files", "*.txt"), ("All files", "*.*")]
            defaultextension = ".json"
        else:
            filetypes = [("Word documents", "*.docx"), ("Text files", "*.txt"), ("All files", "*.*")]
            defaultextension = ".docx"
        
        filename = filedialog.asksaveasfilename(
            title="Save Response As",
            defaultextension=defaultextension,
            filetypes=filetypes
        )
        if filename:
            try:
                if filename.lower().endswith('.json') or (is_json and filename.lower().endswith('.txt')):
                    # Save as JSON
                    if is_json:
                        # Try to format JSON nicely
                        try:
                            json_data = json.loads(response)
                            with open(filename, 'w', encoding='utf-8') as f:
                                json.dump(json_data, f, ensure_ascii=False, indent=2)
                        except json.JSONDecodeError:
                            # If not valid JSON, save as text
                            with open(filename, 'w', encoding='utf-8') as f:
                                f.write(response)
                    else:
                        # Save as text file
                        with open(filename, 'w', encoding='utf-8') as f:
                            f.write(response)
                elif filename.lower().endswith('.docx'):
                    # Save as Word document
                    from docx import Document
                    from datetime import datetime
                    
                    doc = Document()
                    doc.add_heading('PDF Processing Response', 0)
                    
                    # Metadata
                    doc.add_paragraph().add_run('Metadata').bold = True
                    para = doc.add_paragraph()
                    para.add_run(f'Generated: ').bold = True
                    para.add_run(f"{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
                    
                    if self.pdf_path:
                        para = doc.add_paragraph()
                        para.add_run(f'PDF File: ').bold = True
                        para.add_run(os.path.basename(self.pdf_path))
                    
                    para = doc.add_paragraph()
                    para.add_run(f'Model: ').bold = True
                    para.add_run(self.model_var.get())
                    
                    # Response
                    doc.add_paragraph().add_run('Response').bold = True
                    if is_json:
                        # For JSON, add as table or formatted text
                        json_rows = self.parse_json_to_table_data(response)
                        if json_rows:
                            # Add as table
                            table = doc.add_table(rows=len(csv_rows), cols=len(csv_rows[0]))
                            for i, row in enumerate(csv_rows):
                                for j, cell in enumerate(row):
                                    table.rows[i].cells[j].text = str(cell)
                        else:
                            doc.add_paragraph(response)
                    else:
                        response_paragraphs = response.split('\n\n')
                        for para_text in response_paragraphs:
                            if para_text.strip():
                                doc.add_paragraph(para_text.strip())
                    
                    doc.save(filename)
                else:
                    # Save as text file
                    with open(filename, 'w', encoding='utf-8') as f:
                        f.write(response)
                
                messagebox.showinfo("Success", f"Response saved to:\n{filename}")
            except Exception as e:
                messagebox.showerror("Error", f"Failed to save file: {str(e)}")
    
    def open_folder(self, folder_path: str):
        """Open folder in file manager"""
        try:
            import subprocess
            import platform
            
            if platform.system() == "Windows":
                os.startfile(folder_path)
            elif platform.system() == "Darwin":  # macOS
                subprocess.Popen(["open", folder_path])
            else:  # Linux
                subprocess.Popen(["xdg-open", folder_path])
        except Exception as e:
            self.logger.error(f"Error opening folder: {str(e)}")
            messagebox.showerror("Error", f"Could not open folder: {str(e)}")
    
    def json_to_csv_rows(self, json_data: Dict[str, Any]) -> List[List[str]]:
        """
        Convert final_output.json structure to CSV rows format.
        Rows are sorted by page number (Number field).
        
        Args:
            json_data: Dictionary containing 'metadata' and 'rows' keys
            
        Returns:
            List of lists representing CSV rows (first row is headers)
        """
        if not json_data or 'rows' not in json_data:
            return []
        
        rows = json_data['rows']
        if not rows:
            return []
        
        # Sort rows by Part (then by Number for same Part)
        def sort_key(row):
            part = row.get('Part', 0)
            number = row.get('Number', 0)
            try:
                part_num = int(part) if part else 0
            except (ValueError, TypeError):
                part_num = 0
            try:
                number_num = int(number) if number else 0
            except (ValueError, TypeError):
                number_num = 0
            return (part_num, number_num)  # Sort by Part first, then Number
        
        sorted_rows = sorted(rows, key=sort_key)
        
        # CSV headers
        headers = ["Type", "Extraction", "Number", "Part"]
        csv_rows = [headers]
        
        # Convert each row to CSV format
        for row in sorted_rows:
            csv_row = [
                str(row.get('Type', '')),
                str(row.get('Extraction', '')),
                str(row.get('Number', '')),
                str(row.get('Part', ''))
            ]
            csv_rows.append(csv_row)
        
        return csv_rows

    def json_rows_to_csv_rows_generic(self, rows: List[Dict[str, Any]]) -> List[List[str]]:
        """
        Convert a list of JSON row dicts (e.g. second-stage output) to CSV rows.
        Uses all keys as headers (sorted), first row is header.
        """
        if not rows:
            return []
        
        all_keys = set()
        for row in rows:
            if isinstance(row, dict):
                all_keys.update(row.keys())
        
        if not all_keys:
            return []
        
        headers = sorted(list(all_keys))
        csv_rows: List[List[str]] = [headers]
        
        for row in rows:
            if isinstance(row, dict):
                csv_row = [str(row.get(k, "")) for k in headers]
            else:
                csv_row = [str(row)]
            csv_rows.append(csv_row)
        
        return csv_rows
    
    def view_csv_from_json(self):
        """Load final_output.json and display as CSV table"""
        # Try to find final_output.json file
        json_file_path = None
        
        # First, try the stored path
        if self.last_final_output_path and os.path.exists(self.last_final_output_path):
            json_file_path = self.last_final_output_path
        else:
            # Try to find it in current directory
            current_dir = os.getcwd()
            
            # Try exact match first
            exact_path = os.path.join(current_dir, "final_output.json")
            if os.path.exists(exact_path):
                json_file_path = exact_path
            else:
                # Try to find with pattern
                matches = glob.glob(os.path.join(current_dir, "*_final_output.json"))
                if matches:
                    # Use the most recent one
                    json_file_path = max(matches, key=os.path.getmtime)
        
        if not json_file_path:
            # Ask user to select file
            json_file_path = filedialog.askopenfilename(
                title="Select final_output.json file",
                filetypes=[("JSON files", "*.json"), ("All files", "*.*")]
            )
        
        if not json_file_path or not os.path.exists(json_file_path):
            messagebox.showwarning("File Not Found", 
                                 "Could not find final_output.json file.\n\n"
                                 "Please process a PDF first, or select the JSON file manually.")
            return
        
        try:
            # Load JSON file
            with open(json_file_path, 'r', encoding='utf-8') as f:
                json_data = json.load(f)
            
            # Convert to CSV rows
            csv_rows = self.json_to_csv_rows(json_data)
            
            if not csv_rows:
                messagebox.showwarning("Empty Data", "The JSON file contains no rows to display.")
                return
            
            # Create new window for CSV display
            csv_window = ctk.CTkToplevel(self.root)
            csv_window.title(f"CSV View - {os.path.basename(json_file_path)}")
            csv_window.geometry("1200x700")
            csv_window.minsize(800, 500)
            
            # Display CSV table
            self.show_csv_table(csv_window, csv_rows)
            
            # Add info label
            info_label = ctk.CTkLabel(
                csv_window,
                text=f"Displaying {len(csv_rows) - 1} rows from {os.path.basename(json_file_path)}",
                font=ctk.CTkFont(size=12)
            )
            info_label.pack(pady=5)
            
        except json.JSONDecodeError as e:
            messagebox.showerror("JSON Error", f"Failed to parse JSON file:\n{str(e)}")
        except Exception as e:
            messagebox.showerror("Error", f"Failed to load CSV view:\n{str(e)}")
            self.logger.error(f"Error loading CSV view: {str(e)}", exc_info=True)
    
    def run(self):
        """Run the application"""
        self.root.mainloop()


def main():
    """Main entry point"""
    try:
        app = ContentAutomationGUI()
        app.run()
    except Exception as e:
        print(f"Application error: {str(e)}")
        import traceback
        traceback.print_exc()


if __name__ == "__main__":
    main()

